import sys
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import numpy as np
import cv2
import  queue
import time

# 抓包/解包
from packet_capture import PacketCapture
from module_parser import ModuleParser
from network_interface_util import get_network_interfaces


IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}

def levels_pretty(levels: dict) -> str:
    """
    格式化成可读的字符串。
    """
    items = [(a, lv) for a, lv in levels.items() if lv > 0]
    return "，".join([f"{a}: Lv{lv}" for a, lv in items]) if items else "（无有效等级）"

try:
    from ctypes import windll
    windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

BASE_DIR = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))

ATTRS = [
    "抵御魔法","抵御物理","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取",
    "力量加持","敏捷加持","智力加持","特攻伤害加持","精英打击","特攻治疗加持",
    "专精治疗加持","施法专注","攻速专注","暴击专注","幸运专注","极·伤害叠加",
    "极·灵活身法","极·急救措施","极·全队幸暴"
]
NONE_OPTION = "无"
POINT_MIN, POINT_MAX = 1, 10
SLOT_DEFAULT = 4
REQ_KEEP_LIMIT = 100   # 需求最佳
BS_KEEP_LIMIT  = 100   # 战力分推荐
ABS_KEEP_LIMIT = 100   # 绝对最高方案

# 职业对应词条
CLASS_ATTRS = {
    "光盾": {"抵御魔法","抵御物理","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取"},
    "防回": {"抵御魔法","抵御物理","暴击专注","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取"},
    "岩盾": {"抵御魔法","抵御物理","极·绝境守护"},
    "格挡": {"抵御魔法","抵御物理","极·绝境守护","幸运专注"},
    "惩击": {"智力加持","特攻治疗加持","专精治疗加持","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "愈合": {"施法专注","特攻治疗加持","专精治疗加持","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "狂音": {"智力加持","特攻治疗加持","专精治疗加持","攻速专注","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "协奏": {"特攻治疗加持","专精治疗加持","施法专注","极·生命凝聚","极·急救措施","极·全队幸暴","暴击专注"},
    "居合": {"敏捷加持","特攻伤害加持","精英打击","暴击专注","极·伤害叠加","极·灵活身法","极·生命波动"},
    "月刃": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","幸运专注","极·伤害叠加","极·灵活身法","极·生命波动"},
    "重装": {"特攻伤害加持","精英打击","攻速专注","力量加持","极·伤害叠加","极·灵活身法","极·生命波动"},
    "空战": {"特攻伤害加持","精英打击","力量加持","极·伤害叠加","极·灵活身法","极·生命波动","暴击专注","幸运专注"},
    "冰矛": {"智力加持","特攻伤害加持","精英打击","施法专注","暴击专注","幸运专注","极·伤害叠加","极·灵活身法"},
    "射线": {"智力加持","特攻伤害加持","精英打击","施法专注","极·伤害叠加","极·灵活身法"},
    "驭兽": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","极·伤害叠加","极·灵活身法"},
    "驯鹰": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","暴击专注","幸运专注","极·伤害叠加","极·灵活身法"},
}
CLASS_LIST = ["无"] + list(CLASS_ATTRS.keys())

# 多进程批处理评分
from concurrent.futures import ProcessPoolExecutor, as_completed
import itertools, math, heapq, os

# 批大小和并发度可按机器调
MP_BATCH_SIZE  = 50_000
MP_MAX_WORKERS = max(2, (os.cpu_count() or 4) - 1)  # 给系统留1核

def _keep_top(heap, item, limit):
    if len(heap) < limit:
        heapq.heappush(heap, item)
    else:
        if item[0] < heap[0][0]:     # 更优（更小）则替换
            heapq.heapreplace(heap, item)

def _evaluate_batch(mods, idx_batch, targets, prof_attrs,
                    req_keep, abs_keep, bs_keep):
    """
    子进程运行：对一批组合 idx_batch 计算三类榜单的 top 堆。
    返回：(req_top, abs_top, bs_top, processed_count)
    其中每个 top 是 [(key, pack), ...] 的小顶堆（列表）。
    """
    # ---- 把模块内函数/常量拿来用 ----
    summarize_levels_local = summarize_levels
    rank_tuple_local       = rank_tuple
    _battle_key_local      = _battle_key
    format_mod_str_local   = format_mod_str

    req_top, abs_top, bs_top = [], [], []

    for idxs in idx_batch:
        pts = {}
        for i in idxs:
            g = mods[i]
            pts[g["a1"]] = pts.get(g["a1"], 0) + g["p1"]
            pts[g["a2"]] = pts.get(g["a2"], 0) + g["p2"]
            if "a3" in g and "p3" in g:
                pts[g["a3"]] = pts.get(g["a3"], 0) + g["p3"]

        levels, c6, sumlv, waste = summarize_levels_local(pts)
        total_points = sum(pts.values())

        combo_idx  = tuple(idxs)
        combo_mods = tuple(mods[i] for i in idxs)
        pack = (None, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points)

        # 1) 需求榜
        req_key = rank_tuple_local(levels, c6, sumlv, waste, targets, pts_map=pts)
        _keep_top(req_top, (req_key, pack), req_keep)

        # 2) 绝对榜： 6多→等级和高→浪费少→点数和高
        abs_key = (-c6, -sumlv, waste, -total_points)
        _keep_top(abs_top, (abs_key, pack), abs_keep)

        # 3) 职业榜： 6多→职业命中多→再用 req_key
        bs_key = _battle_key_local(levels, req_key, prof_attrs)
        _keep_top(bs_top, (bs_key, pack), bs_keep)

    return (req_top, abs_top, bs_top, len(idx_batch))


def _batched_combinations(n, k, batch_size):
    """
    生成器：把 C(n,k) 组合按批次产出，避免一次性占用超大内存。
    每次 yield 一个 list[tuple[int,...]]，长度约等于 batch_size。
    """
    it = itertools.combinations(range(n), k)
    batch = []
    for comb in it:
        batch.append(comb)
        if len(batch) >= batch_size:
            yield batch
            batch = []
    if batch:
        yield batch

# 等级阈值/评分
def points_to_level(p: int) -> int:
    if p <= 0: return 0
    if p <= 3: return 1
    if p <= 7: return 2
    if p <= 11: return 3
    if p <= 15: return 4
    if p <= 19: return 5
    return 6  # >=20

def format_mod_str(g):
    parts = [f"{g['a1']} {g['p1']}点", f"{g['a2']} {g['p2']}点"]
    if "a3" in g and "p3" in g:
        parts.append(f"{g['a3']} {g['p3']}点")
    return "  +  ".join(parts)

def summarize_levels(points_map):
    lv = {a: points_to_level(v) for a, v in points_map.items()}
    count6 = sum(1 for v in lv.values() if v == 6)
    sumlv = sum(lv.values())
    waste = sum(max(0, v - 20) for v in points_map.values())  # >20 视为浪费点
    return lv, count6, sumlv, waste

def score_for_targets(levels, targets):
    penalty = 0
    for a, req in targets:
        if a and req:
            cur = levels.get(a, 0)
            if cur < req:
                penalty += (req - cur) * 100
    return penalty

def _converted_score(levels: dict) -> int:
    """
    用“2级单位”计算组合强度：
    - Lv6 -> 10
    - Lv5 -> 5
    - Lv2 -> 1
    - 其它等级(0/1/3/4)在这个换算里不计分（保持“只优先考虑 >=5”的语义，
      并能满足 66221 > 5552 的规则）
    """
    n6 = sum(1 for lv in levels.values() if lv == 6)
    n5 = sum(1 for lv in levels.values() if lv == 5)
    n2 = sum(1 for lv in levels.values() if lv == 2)
    return 10 * n6 + 5 * n5 + 1 * n2

def _battle_key(levels: dict, base_rank_key: tuple, prof_attrs: set) -> tuple:
    """
    战力分优先排序键（越小越好）：
    1) 6级数量（多优）
    2) 职业对应词条命中数量（levels>0 且词条在职业表里）（多优）
    3) base_rank_key（你原有的综合排序键：目标惩罚/换算分/等级总和/点数总和等）
    """
    count6 = sum(1 for v in levels.values() if v == 6)
    match = sum(1 for a, lv in levels.items() if lv > 0 and a in prof_attrs) if prof_attrs else 0
    # 想“多者优先”，升序里用负号
    return (-count6, -match, base_rank_key)

def rank_tuple(levels, count6, sumlv, waste, targets, pts_map=None):
    """
    排序键（越小越好）：
    1) 目标惩罚（差1级=+100） —— 有目标时才会产生差异
    2) 换算分（6=10, 5=5, 2=1；越大越好）
    3) 再看“6 的数量”更多者优先
    4) 再看“≥5 的具体向量”字典序（保证 6>5，不把5当6）
    5) 再看 等级总和（越大越好）
    6) 再看 点数总和（越大越好；包含>20的溢出）
    """
    penalty = score_for_targets(levels, targets)
    total_points = sum(pts_map.values()) if pts_map else 0

    # 换算分（越大越好 → 取负用于“升序排序”）
    conv = _converted_score(levels)

    # 更细的高等级比较（不把5当6）
    highs = sorted((lv for lv in levels.values() if lv >= 5), reverse=True)
    n6 = sum(1 for lv in highs if lv == 6)
    high_vec_key = tuple(-lv for lv in highs)  # 降序比较 → 升序排序时取负

    # 等级总和、点数总和（都取负表示“大者优先”）
    return (
        penalty,          # 目标惩罚（小优）
        -conv,            # 换算分（大优）
        -n6,              # 6 的数量（多优）
        high_vec_key,     # ≥5 的向量字典序（保证 6>5）
        -sumlv,           # 等级总和（大优）
        -total_points,    # 点数总和（大优；含溢出）
    )

class SimpleListDialog(tk.Toplevel):
    """简单单选列表弹窗：支持搜索、单选、回车确认、双击选择。"""
    def __init__(self, master, title, options, current=None):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.result = None
        self._opts_all = list(options)

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # 搜索框
        top = ttk.Frame(frm)
        top.pack(fill=tk.X)
        ttk.Label(top, text="搜索：").pack(side=tk.LEFT)
        self.var_q = tk.StringVar()
        ent = ttk.Entry(top, textvariable=self.var_q, width=24)
        ent.pack(side=tk.LEFT, padx=(6,0))
        ent.bind("<KeyRelease>", self._on_search)

        # 列表
        self.lb = tk.Listbox(frm, height=12, exportselection=False)
        self.lb.pack(fill=tk.BOTH, expand=True, pady=8)
        self.lb.bind("<Double-Button-1>", lambda e: self._confirm())
        self.lb.bind("<Return>", lambda e: self._confirm())
        self._reload_list()

        # 预选
        if current in self._opts_all:
            idx = self._opts_all.index(current)
            self.lb.selection_clear(0, tk.END)
            self.lb.selection_set(idx)
            self.lb.see(idx)

        # 底部按钮
        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X)

        self.update_idletasks()
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def _reload_list(self, filtered=None):
        data = filtered if filtered is not None else self._opts_all
        self.lb.delete(0, tk.END)
        for it in data:
            self.lb.insert(tk.END, it)

    def _on_search(self, *_):
        q = self.var_q.get().strip().lower()
        if not q:
            self._reload_list()
            return
        filt = [o for o in self._opts_all if q in o.lower()]
        self._reload_list(filt)

    def _confirm(self):
        sel = self.lb.curselection()
        if not sel:
            return
        self.result = self.lb.get(sel[0])
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()


# ========= 弹窗：词条选择器 =========
class AttributePicker(tk.Toplevel):
    """
    简单弹窗：搜索框 + 列表 + 确定/取消。
    - 支持禁用项（disabled_set），不可选；
    - 双击或选中后点“确定”返回；
    - on_ok(name) 回调返回所选字符串，或 None 表示取消。
    """
    def __init__(self, master, all_attrs, disabled_set=None, title="选择词条", on_ok=None, init_text=""):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

        self.all_attrs = list(all_attrs)
        self.disabled = set(disabled_set or [])
        self.on_ok = on_ok

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        top = ttk.Frame(frm)
        top.pack(fill=tk.X)
        ttk.Label(top, text="搜索：").pack(side=tk.LEFT)
        self.var_q = tk.StringVar(value=init_text)
        ent = ttk.Entry(top, textvariable=self.var_q, width=28)
        ent.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(4,0))
        ent.bind("<KeyRelease>", self._on_filter)

        mid = ttk.Frame(frm)
        mid.pack(fill=tk.BOTH, expand=True, pady=(8,6))
        self.lb = tk.Listbox(mid, height=12)
        self.lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self.lb.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.lb.configure(yscrollcommand=sb.set)
        self.lb.bind("<Double-Button-1>", self._on_dbl)

        bot = ttk.Frame(frm)
        bot.pack(fill=tk.X)

        self._reload_list()
        self._center_to(master)

    def _center_to(self, master):
        try:
            self.update_idletasks()
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def _on_filter(self, e=None):
        self._reload_list()

    def _reload_list(self):
        q = self.var_q.get().strip().lower()
        cand = [a for a in self.all_attrs if (q in a.lower())] if q else list(self.all_attrs)
        cand.sort()
        try:
            from mod import NONE_OPTION
        except Exception:
            NONE_OPTION = "无"
        if NONE_OPTION in cand:
            cand = [NONE_OPTION] + [a for a in cand if a != NONE_OPTION]

        self.lb.delete(0, tk.END)
        for a in cand:
            if a in self.disabled:
                self.lb.insert(tk.END, f"× {a}")
            else:
                self.lb.insert(tk.END, a)

    def _pick_value(self):
        cur = self.lb.curselection()
        if not cur:
            return None
        val = self.lb.get(cur[0])
        if val.startswith("× "):
            val = val[2:]
            from tkinter import messagebox
            messagebox.showwarning("不可选择", f"“{val}”已在另一栏被占用/冲突。")
            return None
        return val

    def _on_dbl(self, e=None):
        v = self._pick_value()
        if v and callable(self.on_ok):
            self.on_ok(v)
            self.destroy()

    def _on_ok(self):
        v = self._pick_value()
        if v and callable(self.on_ok):
            self.on_ok(v)
            self.destroy()

    def _on_cancel(self):
        if callable(self.on_ok):
            self.on_ok(None)
        self.destroy()


    """
    勾选框多选弹窗：
    - all_attrs: 全部可选词条（list[str]）
    - disabled_set: 需要禁用的词条（set[str]），用于目标/排除互斥
    - preselected: 预选中的词条（list[str]）
    - max_select: 最多可勾选数量（int）
    - on_ok: 回调 on_ok(selected_list)；取消时返回 None
    """
    def __init__(self, master, all_attrs, disabled_set=None, preselected=None, max_select=5, title="选择词条（多选）", on_ok=None):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

        self.all_attrs   = list(all_attrs)
        self.disabled    = set(disabled_set or [])
        self.on_ok       = on_ok
        self.max_select  = int(max_select or 0)
        self.vars_by_attr = {}  # attr -> BooleanVar

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # 顶部：搜索 & 计数
        top = ttk.Frame(frm); top.pack(fill=tk.X)
        ttk.Label(top, text="搜索：").pack(side=tk.LEFT)
        self.var_q = tk.StringVar()
        ent = ttk.Entry(top, textvariable=self.var_q, width=24)
        ent.pack(side=tk.LEFT, padx=(6,10))
        self.lbl_cnt = ttk.Label(top, text=f"已选 0 / {self.max_select}")
        self.lbl_cnt.pack(side=tk.LEFT)

        # 中部：滚动区域 + 勾选框
        mid = ttk.Frame(frm); mid.pack(fill=tk.BOTH, expand=True, pady=8)
        canvas = tk.Canvas(mid, width=360, height=280, highlightthickness=0)
        vsb = ttk.Scrollbar(mid, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.chk_wrap = ttk.Frame(canvas)
        canvas.create_window((0,0), window=self.chk_wrap, anchor="nw")

        def _on_config(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
        self.chk_wrap.bind("<Configure>", _on_config)

        # 底部：按钮
        bot = ttk.Frame(frm); bot.pack(fill=tk.X, pady=(2,0))
        ttk.Button(bot, text="清空", command=self._clear_all).pack(side=tk.LEFT)
        ttk.Button(bot, text="确定", command=self._ok).pack(side=tk.RIGHT, padx=(6,0))
        ttk.Button(bot, text="取消", command=self._on_cancel).pack(side=tk.RIGHT)

        # 事件
        ent.bind("<KeyRelease>", lambda e: self._reload_checks())

        # 初始化
        self._reload_checks(preselected=set(preselected or []))
        self._center_to(master)

    def _center_to(self, master):
        try:
            self.update_idletasks()
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def _reload_checks(self, preselected=None):
        # 根据搜索重建勾选列表；保留已存在的 var
        q = (self.var_q.get() or "").strip().lower()
        for w in list(self.chk_wrap.winfo_children()):
            w.destroy()

        # “无”置顶（如果存在），其余按字典序
        items = list(self.all_attrs)
        if "无" in items:
            items = ["无"] + [a for a in items if a != "无"]
        else:
            items.sort()

        for a in items:
            if q and (q not in a.lower()):
                continue
            var = self.vars_by_attr.get(a) or tk.BooleanVar(value=False)
            self.vars_by_attr[a] = var
            if preselected and a in preselected:
                var.set(True)
            state = ("disabled" if a in self.disabled else "normal")
            chk = ttk.Checkbutton(self.chk_wrap, text=a, variable=var, state=state,
                                  command=self._enforce_limit)
            chk.pack(anchor="w", pady=2)

        self._enforce_limit()

    def _selected_list(self):
        return [a for a, v in self.vars_by_attr.items() if v.get()]

    def _enforce_limit(self):
        sel = self._selected_list()
        n = len(sel)
        self.lbl_cnt.config(text=f"已选 {n} / {self.max_select}")
        # 达上限时，除已选和已禁用外，其余全部禁用
        lock_others = (self.max_select > 0 and n >= self.max_select)
        for a, var in self.vars_by_attr.items():
            chk_state = "normal"
            if a in self.disabled:
                chk_state = "disabled"
            elif lock_others and (a not in sel):
                chk_state = "disabled"
            # 找到对应的 Checkbutton 设置状态
            # 这里通过遍历子控件绑定的 text 来定位
            for w in self.chk_wrap.winfo_children():
                try:
                    if isinstance(w, ttk.Checkbutton) and w.cget("text") == a:
                        w.state(["disabled"] if chk_state == "disabled" else ["!disabled"])
                        break
                except Exception:
                    pass

    def _clear_all(self):
        for var in self.vars_by_attr.values():
            var.set(False)
        self._enforce_limit()

    def _ok(self):
        if callable(self.on_ok):
            self.on_ok(self._selected_list())
        self.destroy()

    def _on_cancel(self):
        if callable(self.on_ok):
            self.on_ok(None)
        self.destroy()

class ClickSelectEntry(ttk.Frame):
    """
    无下拉、无按钮：只读输入框，点击后弹出选择窗口。
    """
    def __init__(self, master, options, placeholder="", width=16, on_select=None):
        super().__init__(master)
        self.options = list(options)
        self.placeholder = placeholder
        self.on_select = on_select
        self._popup = None  # 防止重复弹出

        self.var = tk.StringVar(value=placeholder)
        self.entry = ttk.Entry(self, textvariable=self.var, width=width, state="readonly")
        self.entry.pack(fill="x", expand=True)
        self.entry.bind("<Button-1>", self._open_popup)

    def get(self):
        v = (self.var.get() or "").strip()
        return "" if v == self.placeholder else v

    def set(self, v: str):
        self.var.set(v or self.placeholder)

    def _open_popup(self, event=None):
        if not self.options or (self._popup and self._popup.winfo_exists()):
            return
        top = tk.Toplevel(self)
        self._popup = top
        top.title("请选择")
        top.transient(self.winfo_toplevel())
        top.grab_set()
        top.resizable(False, False)
        top.protocol("WM_DELETE_WINDOW", lambda: self._close_popup(top))

        frm = ttk.Frame(top, padding=8)
        frm.pack(fill="both", expand=True)

        sv = tk.StringVar()
        ent_search = ttk.Entry(frm, textvariable=sv)
        ent_search.pack(fill="x", pady=(0,6))
        ent_search.focus_set()

        lb = tk.Listbox(frm, height=12)
        lb.pack(fill="both", expand=True)

        def _refresh():
            kw = sv.get().strip().lower()
            lb.delete(0, tk.END)
            for it in self.options:
                if not kw or kw in it.lower():
                    lb.insert(tk.END, it)
        _refresh()

        def _ok():
            sel = lb.curselection()
            if sel:
                val = lb.get(sel[0])
                if val == "无":
                    self.var.set("")
                else:
                    self.var.set(val)
                if callable(self.on_select):
                    self.on_select(val if val != "无" else "")
            self._close_popup(top)

        def _cancel():
            self._close_popup(top)

        lb.bind("<Double-Button-1>", lambda e: _ok())
        lb.bind("<Return>", lambda e: _ok())

        btns = ttk.Frame(frm)
        btns.pack(fill="x", pady=(6,0))

        sv.trace_add("write", lambda *_: _refresh())

        top.update_idletasks()
        try:
            x = self.winfo_rootx() + (self.winfo_width()-top.winfo_width())//2
            y = self.winfo_rooty() + self.winfo_height() + 6
            top.geometry(f"+{x}+{y}")
        except Exception:
            pass

    def _close_popup(self, top):
        try:
            top.grab_release()
        except Exception:
            pass
        try:
            top.destroy()
        except Exception:
            pass
        self._popup = None


# ========= 组合控件：只读显示=========
class AttrPickerField(ttk.Frame):
    def __init__(self, master, values, on_change=None, width=24, placeholder="",
                 multi_mode=False, max_select=1, get_pre=None, get_disabled=None,
                 apply_selected=None, picker_title="选择词条（多选）"):
        super().__init__(master)
        self.values = list(values)
        self.disabled = set()
        self.on_change = on_change
        self.placeholder = placeholder or ""

        # 多选相关
        self.multi_mode = bool(multi_mode)
        self.max_select = int(max_select)
        self.get_pre = get_pre
        self.get_disabled = get_disabled
        self.apply_selected = apply_selected
        self.picker_title = picker_title

        self.var = tk.StringVar(value=self.placeholder)
        self.ent = ttk.Entry(self, textvariable=self.var, width=width, state="readonly")
        self.ent.pack(side=tk.LEFT, fill="x", expand=True)
        self.ent.bind("<Button-1>", self._open_picker)

    def _open_picker(self, *_):
        # 多选模式：打开 MultiAttrPicker
        if self.multi_mode:
            pre = self.get_pre() if callable(self.get_pre) else []
            disabled = self.get_disabled() if callable(self.get_disabled) else set()
            # 这里不提供“无”，因此从 values 里排除“无”
            opts = [a for a in self.values if a != "无"]
            MultiAttrPicker(
                self.winfo_toplevel(),
                all_attrs=opts,
                disabled_set=disabled,
                preselected=pre,
                max_select=self.max_select,
                title=self.picker_title or "选择词条（多选）",
                on_ok=lambda selected: (self.apply_selected(selected) if (selected is not None and callable(self.apply_selected)) else None)
            )
            return

        # 单选模式：维持旧逻辑（用于“添加模组”的词条1/2等）
        init = self.get_value() or self.placeholder
        AttributePicker(
            self.winfo_toplevel(),
            self.values,
            disabled_set=self.disabled,
            title="选择词条",
            on_ok=self._on_selected,
            init_text=""  # 默认搜索框留空，避免必须清空才能选（你之前的诉求）
        )

    def _on_selected(self, name):
        if name is None:
            return
        if name == "无":
            self.var.set(self.placeholder)
        else:
            self.var.set(name)
        if callable(self.on_change):
            self.on_change(self)

    def set_disabled(self, names:set):
        self.disabled = set(names or [])
        cur = self.get_value()
        if cur and cur in self.disabled:
            self.var.set(self.placeholder)
            if callable(self.on_change):
                self.on_change(self)

    def get_value(self):
        v = (self.var.get() or "").strip()
        return "" if v == self.placeholder else v

class MultiAttrPicker(tk.Toplevel):
    """
    多选勾选框弹窗（无滚动条、每行4个）：
    - all_attrs: 候选词条（调用方请不要传入“无”）
    - disabled_set: 需要禁用的词条（互斥用）
    - preselected: 预选词条
    - max_select: 最多可勾选数量（例如 5）
    - on_ok: on_ok(selected_list)；取消时回调 on_ok(None)
    """
    def __init__(self, master, all_attrs, disabled_set=None, preselected=None,
                 max_select=5, title="选择词条（多选）", on_ok=None):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)   # 固定大小；如果想允许用户拉伸可改 True, True
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

        # 数据
        self.all_attrs = sorted([a for a in all_attrs if a != "无"])
        self.disabled  = set(disabled_set or [])
        self.on_ok     = on_ok
        self.max_select= int(max_select or 0)

        # 状态
        self.vars_by_attr: dict[str, tk.BooleanVar] = {}
        self.chk_widgets:  dict[str, ttk.Checkbutton] = {}

        # 布局
        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # 顶部：计数
        top = ttk.Frame(frm); top.pack(fill=tk.X)
        self.lbl_cnt = ttk.Label(top, text=f"已选 0 / {self.max_select}")
        self.lbl_cnt.pack(side=tk.LEFT)

        # 中部：直接网格排版（无滚动条）
        self.chk_wrap = ttk.Frame(frm)
        self.chk_wrap.pack(fill=tk.BOTH, expand=True, pady=8)

        self._cols = 4  # 每行 4 个
        for c in range(self._cols):
            self.chk_wrap.grid_columnconfigure(c, weight=1)

        # 底部：按钮
        bot = ttk.Frame(frm); bot.pack(fill=tk.X, pady=(2, 0))
        ttk.Button(bot, text="清空", command=self._clear_all).pack(side=tk.LEFT)
        ttk.Button(bot, text="确定", command=self._ok).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(bot, text="取消", command=self._on_cancel).pack(side=tk.RIGHT)

        # 生成勾选项并居中
        self._build_checks(preselected=set(preselected or []))
        self._center_to(master)

    # ---------- 内部方法 ----------
    def _center_to(self, master):
        try:
            self.update_idletasks()
            x = master.winfo_rootx() + (master.winfo_width() - self.winfo_width()) // 2
            y = master.winfo_rooty() + (master.winfo_height() - self.winfo_height()) // 2
            self.geometry(f"+{max(0, x)}+{max(0, y)}")
        except Exception:
            pass

    def _build_checks(self, preselected:set[str]):
        # 清空旧控件
        for w in list(self.chk_wrap.winfo_children()):
            w.destroy()
        self.vars_by_attr.clear()
        self.chk_widgets.clear()

        # 逐项生成（4列网格）
        for i, a in enumerate(self.all_attrs):
            r, c = divmod(i, self._cols)
            var = tk.BooleanVar(value=(a in preselected))
            self.vars_by_attr[a] = var
            state = ("disabled" if a in self.disabled else "normal")
            chk = ttk.Checkbutton(self.chk_wrap, text=a, variable=var,
                                  state=state, command=self._enforce_limit)
            chk.grid(row=r, column=c, sticky="w", padx=8, pady=6)
            self.chk_widgets[a] = chk

        self._enforce_limit()

    def _selected_list(self):
        return [a for a, v in self.vars_by_attr.items() if v.get()]

    def _enforce_limit(self):
        sel = self._selected_list()
        n = len(sel)
        self.lbl_cnt.config(text=f"已选 {n} / {self.max_select}")

        lock_others = (self.max_select > 0 and n >= self.max_select)
        for a, chk in self.chk_widgets.items():
            if a in self.disabled:
                chk.state(["disabled"])
            elif lock_others and (a not in sel):
                chk.state(["disabled"])
            else:
                chk.state(["!disabled"])

    def _clear_all(self):
        for var in self.vars_by_attr.values():
            var.set(False)
        self._enforce_limit()

    def _ok(self):
        if callable(self.on_ok):
            self.on_ok(self._selected_list())
        self.destroy()

    def _on_cancel(self):
        if callable(self.on_ok):
            self.on_ok(None)
        self.destroy()


# ========= 进度弹窗 =========
class _Progress(tk.Toplevel):
    def __init__(self, master, total, on_cancel):
        super().__init__(master)
        self.title("正在计算…")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", on_cancel)
        self.grab_set()  # 模态
        self.var = tk.StringVar(value="准备中…")
        ttk.Label(self, textvariable=self.var).pack(padx=14, pady=(12,6))
        self.pb = ttk.Progressbar(self, length=360, mode="determinate", maximum=max(1, total))
        self.pb.pack(padx=14, pady=6)
        ttk.Button(self, text="取消", command=on_cancel).pack(pady=(6,12))
        self.update_idletasks()
        # 居中
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def set_progress(self, done, total):
        self.pb["value"] = min(done, total)
        self.var.set(f"已处理 {done:,} / {total:,} 个组合…")

# ========= 主应用 =========
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("模组词条配点计算器   by椛椿")
        self.geometry("1650x1000")
        self.minsize(1080, 680)
        self.mods = []  # 每项: {"a1":..., "p1":int, "a2":..., "p2":int}

        # 记录互斥输入框引用
        self.targets = []
        self.target_entries = []
        self.exclude_entries = []

        # 异步计算通信
        self._q = queue.Queue()
        self._worker = None
        self._cancel = False

        self.build_ui()
        self.var_mod_title = tk.StringVar(value="模组词条 (0)")

        # 抓包/解包状态
        self.capture = None            # PacketCapture 实例
        self.module_parser = ModuleParser()
        self._iface_list = []          # [(显示名, 实际名)]
        self._last_mod_time = 0        # 最近一次收到模组时间戳
        self._idle_timeout = 60       # “无新模组自动停止”的秒数，可自行调整
        self._auto_stop_job = None     # after 定时器句柄

    def export_solutions_to_docx(self, which: str) -> None:
        """
        导出方案为 Word（.docx）。
        which ∈ {"req", "abs", "bs"} 分别对应：需求最佳 / 最高等级 / 战力分推荐。
        依赖：python-docx  (pip install python-docx)
        """
        # 1) 依赖检查
        try:
            from docx import Document
        except Exception:
            from tkinter import messagebox
            messagebox.showerror("缺少依赖", "需要安装 python-docx 才能导出 Word：\n\npip install python-docx")
            return

        # 2) 选择数据源
        if which == "req":
            data = getattr(self, "_last_req_list", [])
            title = "需求最佳方案（前100）"
        elif which == "abs":
            data = getattr(self, "_last_abs_list", [])
            title = "最高等级方案（前100）"
        elif which == "bs":
            data = getattr(self, "_last_bs_list", [])
            title = "战力分推荐方案（前100）"
        else:
            from tkinter import messagebox
            messagebox.showerror("错误", f"未知导出类型：{which}")
            return

        if not data:
            from tkinter import messagebox
            messagebox.showwarning("提示", "当前没有可导出的结果。请先计算。")
            return

        # 3) 生成文档
        doc = Document()
        doc.add_heading(title, level=1)

        def levels_pretty(levels: dict) -> str:
            items = [(a, lv) for a, lv in levels.items() if lv > 0]
            items.sort(key=lambda x: (-x[1], x[0]))
            return "，".join([f"{a}: Lv{lv}" for a, lv in items]) if items else "（无有效等级）"

        def lv6_attrs(levels: dict) -> str:
            arr = [a for a, lv in levels.items() if lv == 6]
            return "，".join(arr) if arr else "（无）"

        for k, pack in enumerate(data, start=1):
            # 统一解包： (rank_key/None, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points)
            _, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points = pack

            doc.add_heading(f"方案 #{k}", level=2)

            # 使用模组：逐行显示
            for g in combo_mods:
                # 这里沿用你现有的 format_mod_str(g)
                doc.add_paragraph(str(format_mod_str(g)))

            # 概览与统计
            doc.add_paragraph("等级概览：" + levels_pretty(levels))
            doc.add_paragraph("6级词条：" + lv6_attrs(levels))
            doc.add_paragraph(f"统计：6级数量={c6}，等级总和={sumlv}，浪费点={waste}，点数总和={total_points}")

        # 4) 保存对话框
        from tkinter import filedialog, messagebox
        path = filedialog.asksaveasfilename(
            title="导出为 Word",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile=title.replace("（前100）", "").strip() + ".docx"
        )
        if not path:
            return

        # 5) 保存 & 完成提示
        try:
            doc.save(path)
        except Exception as e:
            messagebox.showerror("保存失败", f"无法写入文件：\n{e}")
            return

        messagebox.showinfo("完成", f"已导出到：\n{path}")

    def _compute_mp(self, mods, slots, targets, total_combos, prof_attrs):
        """
        多进程协调器（跑在单独的线程里）：
        - 逐批把组合分发给进程池
        - 合并各子进程返回的三类 top 堆
        - 通过 self._q 持续汇报进度 & 最终结果
        """
        import heapq
        processed = 0

        # 全局 top 堆（与原单线程版一致）
        req_top, abs_top, bs_top = [], [], []

        def keep_top(heap, item, limit):
            if len(heap) < limit:
                heapq.heappush(heap, item)
            else:
                if item[0] < heap[0][0]:
                    heapq.heapreplace(heap, item)

        # 预先取到三个榜单的保留上限（与你的常量一致）
        req_keep = REQ_KEEP_LIMIT
        abs_keep = ABS_KEEP_LIMIT
        bs_keep  = BS_KEEP_LIMIT

        # 生成批次（不会一次性把所有组合存内存）
        m = len(mods)
        batches = _batched_combinations(m, slots, MP_BATCH_SIZE)

        # 进程池并发执行；控制在飞任务数量
        in_flight = []
        finished  = False

        try:
            import multiprocessing as mp
            ctx = mp.get_context("spawn")
            with ProcessPoolExecutor(max_workers=MP_MAX_WORKERS, mp_context=ctx) as pool:
                # 先铺一批任务
                MAX_IN_FLIGHT = MP_MAX_WORKERS * 3
                for batch in batches:
                    if self._cancel:
                        break
                    fut = pool.submit(_evaluate_batch, mods, batch, targets, prof_attrs,
                                    req_keep, abs_keep, bs_keep)
                    in_flight.append((fut, len(batch)))

                    # 控制在飞数
                    if len(in_flight) >= MAX_IN_FLIGHT:
                        # 等一个完成再继续提交
                        done_fut = None
                        for fut_i, batch_len in list(in_flight):
                            if fut_i.done():
                                done_fut = (fut_i, batch_len)
                                in_flight.remove((fut_i, batch_len))
                                break
                        if done_fut is None:
                            # 阻塞等待任意一个
                            fut_i, batch_len = in_flight.pop(0)
                            res = fut_i.result()
                            # 合并
                            _req, _abs, _bs, cnt = res
                            for it in _req: keep_top(req_top, it, req_keep)
                            for it in _abs: keep_top(abs_top, it, abs_keep)
                            for it in _bs:  keep_top(bs_top,  it, bs_keep)
                            processed += cnt
                            self._q.put(("progress", (processed, total_combos)))

                # 把余下 in_flight 都收掉
                for fut_i, batch_len in in_flight:
                    if self._cancel:
                        break
                    res = fut_i.result()
                    _req, _abs, _bs, cnt = res
                    for it in _req: keep_top(req_top, it, req_keep)
                    for it in _abs: keep_top(abs_top, it, abs_keep)
                    for it in _bs:  keep_top(bs_top,  it, bs_keep)
                    processed += cnt
                    self._q.put(("progress", (processed, total_combos)))

            finished = True

        except Exception:
            # 出错时也尽量返回已经拿到的结果
            finished = True

        # 用户点“取消”
        if self._cancel:
            self._q.put(("cancelled", None))
            return

        if finished:
            # 小顶堆 → 升序列表（与原渲染接口一致：只要 pack 列表）
            req_list = [x[1] for x in sorted(req_top, key=lambda t: t[0])]
            abs_list = [x[1] for x in sorted(abs_top, key=lambda t: t[0])]
            bs_list  = [x[1] for x in sorted(bs_top,  key=lambda t: t[0])]
            max_c6   = max((p[5] for p in abs_list), default=0)
            self._q.put(("done", (req_list, abs_list, bs_list, max_c6, processed, total_combos)))


    def delete_low_score_mods(self):
        import tkinter.simpledialog as sd
        # 输入阈值（<=14）
        x = sd.askinteger("删除低级模组", "请输入点数阈值（最大14）：", minvalue=0, maxvalue=14)
        if x is None:
            return
        # 统计拟删除
        cand = [i for i, g in enumerate(self.mods) if (g["p1"] + g["p2"]) <= x]
        if not cand:
            messagebox.showinfo("提示", f"没有点数和 < {x} 的模组。")
            return
        if not messagebox.askyesno("确认", f"将删除 {len(cand)} 个模组（点数和 < {x}）。是否继续？"):
            return
        # 从后往前删，避免索引位移
        for i in sorted(cand, reverse=True):
            del self.mods[i]
        # 重建列表
        for item in self.tree.get_children():
            self.tree.delete(item)
        for g in self.mods:
            self.tree.insert("", tk.END, values=(format_mod_str(g),))
        messagebox.showinfo("完成", f"已删除 {len(cand)} 个模组。")
        self._update_mod_counters()

    # 回填/预选/互斥集合
    #职业
    def _get_pre_job(self):
        try:
            v = self.job_field.get_value()
        except Exception:
            v = ""
        return [v] if v else []

    def _apply_selected_job(self, selected):
        if selected is None:
            return
        name = selected[0] if selected else ""
        if name:
            self.job_field.var.set(name)
        else:
            self.job_field.var.set(self.job_field.placeholder or "")
        if hasattr(self, "_entry_changed") and callable(self._entry_changed):
            self._entry_changed(self.job_field)

    #词条
    def _current_targets_set(self):
        s = set()
        for ent, _ in self.targets:
            v = ent.get_value()
            if v: s.add(v)
        return s

    def _current_excludes_set(self):
        s = set()
        for ent in self.exclude_entries:
            v = ent.get_value()
            if v: s.add(v)
        return s

    def _get_preselected_targets(self):
        return [ent.get_value() for ent, _ in self.targets if ent.get_value()]

    def _get_preselected_excludes(self):
        return [ent.get_value() for ent in self.exclude_entries if ent.get_value()]

    def _apply_selected_targets(self, selected):
        if selected is None: return
        # 去重并保序
        selected = list(dict.fromkeys(selected))
        max_sel = len(self.targets)
        selected = selected[:max_sel]
        for i, (ent, sp) in enumerate(self.targets):
            ent.var.set(selected[i] if i < len(selected) else (ent.placeholder or ""))
        self._refresh_mutex_states(False)

    def _apply_selected_excludes(self, selected):
        if selected is None: return
        selected = list(dict.fromkeys(selected))
        max_sel = len(self.exclude_entries)
        selected = selected[:max_sel]
        for i, ent in enumerate(self.exclude_entries):
            ent.var.set(selected[i] if i < len(selected) else (ent.placeholder or ""))
        self._refresh_mutex_states(False)


    # 互斥联动：当任一目标/排除改动时调用
    def _entry_changed(self, _entry):
        self._refresh_mutex_states(show_prompt=True)

    # 计算当前已选目标/排除集合，刷新禁用项，并对冲突给出提示/清空
    def _refresh_mutex_states(self, show_prompt=False):
        targets_set = set()
        for ent, sp in self.targets:
            v = ent.get_value()
            if v and v in ATTRS:
                targets_set.add(v)
        excludes_set = set()
        for ent in self.exclude_entries:
            v = ent.get_value()
            if v and v in ATTRS:
                excludes_set.add(v)

        # 互斥：目标禁用“排除已选”，排除禁用“目标已选”
        for ent, _ in self.targets:
            ent.set_disabled(excludes_set)
        for ent in self.exclude_entries:
            ent.set_disabled(targets_set)

        # 冲突时清空排除侧的冲突项
        overlap = targets_set & excludes_set
        if overlap:
            if show_prompt:
                messagebox.showwarning("冲突", f"目标词条与排除词条不能相同：{', '.join(sorted(overlap))}")
            for ent in self.exclude_entries:
                if ent.get_value() in overlap:
                    ent.var.set("")
            self.after(10, lambda: self._refresh_mutex_states(False))

    def _update_mod_counters(self):
        cnt = len(self.mods)
        # 刷新 LabelFrame 标题
        if hasattr(self, "listf"):
            self.listf.configure(text=f"已添加的模组 ({cnt})")
        # 刷新 Treeview 列表头
        if hasattr(self, "tree"):
            self.tree.heading("desc", text=f"模组词条 ({cnt})")

    def build_ui(self):
        # 外层左右两栏
        root = ttk.Frame(self, padding=12)
        root.pack(fill=tk.BOTH, expand=True)

        left = ttk.Frame(root)
        right = ttk.Frame(root)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0,8))
        right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # —— 工具条（同一行；仅调整顺序与标签宽度/边距） ——
        tool = ttk.Frame(left)  
        tool.pack(fill=tk.X, pady=(0,6))
        LABEL_W = 6  # 职业/网卡标签统一宽度

        ttk.Label(tool, text="网卡：", width=LABEL_W).pack(side=tk.LEFT, padx=(6,2))
        self.var_iface = tk.StringVar(value="")
        self.cb_iface = ttk.Combobox(tool, textvariable=self.var_iface, width=34, state="readonly")
        self.cb_iface.pack(side=tk.LEFT, padx=(0,6))
        ttk.Button(tool, text="刷新网卡", command=self._refresh_interfaces).pack(side=tk.LEFT)
        ttk.Button(tool, text="开始抓包", command=self._start_sniff).pack(side=tk.LEFT, padx=(10,6))
        self.btn_stop = ttk.Button(tool, text="停止抓包", command=self._stop_sniff, state="disabled")
        self.btn_stop.pack(side=tk.LEFT, padx=(0,10))
        # 首次进入自动刷新一次网卡
        self._refresh_interfaces()

        # 左侧：职业/期望/排除/计算
        #职业
        jobf = ttk.Frame(left)
        job_options = ["无"] + [c for c in CLASS_LIST if c != "无"]
        self.jobs_selected = []
        jobf.pack(fill=tk.X, pady=(0,6))
        ttk.Label(jobf, text="职业：", width=LABEL_W).pack(side=tk.LEFT, padx=(6,2))
        self.job_field = AttrPickerField(
            jobf,
            job_options,
            on_change=self._entry_changed,
            width=12,
            multi_mode=True,
            max_select=1,
            get_pre=self._get_pre_job,
            get_disabled=lambda: set(),
            apply_selected=self._apply_selected_job,
            picker_title="选择职业）"
        )
        self.job_field.pack(side=tk.LEFT, padx=4)

        # 目标词条
        targetf = ttk.LabelFrame(left, text="目标词条和等级（可不填）")
        targetf.pack(fill=tk.X, pady=(6,6))

        cols = ttk.Frame(targetf)
        cols.pack(fill=tk.BOTH, expand=True, padx=6, pady=4)

        col_left = ttk.Frame(cols)
        col_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        col_right = ttk.Frame(cols)
        col_right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(16,0))  # 列间距

        # 左列：目标 1~3
        for i in range(3):
            rowf = ttk.Frame(col_left)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"目标{i+1}").pack(side=tk.LEFT)
            ent = AttrPickerField(
                rowf, ATTRS, on_change=self._entry_changed, width=12,
                multi_mode=True,
                max_select=5,
                get_pre=self._get_preselected_targets, 
                get_disabled=self._current_excludes_set,
                apply_selected=self._apply_selected_targets,
                picker_title="选择目标词条"
            )
            ent.pack(side=tk.LEFT, padx=(6,8))
            ttk.Label(rowf, text="期望等级").pack(side=tk.LEFT)
            sp = tk.Spinbox(rowf, from_=0, to=6, width=5)
            sp.delete(0, tk.END); sp.insert(0, "0")
            sp.pack(side=tk.LEFT, padx=6)
            self.targets.append((ent, sp))
            self.target_entries.append(ent)

        # 右列：目标 4~5
        for i in range(3, 5):
            rowf = ttk.Frame(col_right)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"目标{i+1}").pack(side=tk.LEFT)
            ent = AttrPickerField(
                rowf, ATTRS, on_change=self._entry_changed, width=12,
                multi_mode=True,
                max_select=5,
                get_pre=self._get_preselected_targets,
                get_disabled=self._current_excludes_set,
                apply_selected=self._apply_selected_targets,
                picker_title="选择目标词条"
            )
            ent.pack(side=tk.LEFT, padx=(0,10))
            ttk.Label(rowf, text="期望等级").pack(side=tk.LEFT)
            sp = tk.Spinbox(rowf, from_=0, to=6, width=5)
            sp.delete(0, tk.END); sp.insert(0, "0")
            sp.pack(side=tk.LEFT, padx=6)
            self.targets.append((ent, sp))
            self.target_entries.append(ent)

        # 排除词条
        exf = ttk.LabelFrame(left, text="排除词条（可不填）")
        exf.pack(fill=tk.X, pady=(6,6))

        self.exclude_entries = []

        ex_cols = ttk.Frame(exf)
        ex_cols.pack(fill=tk.BOTH, expand=True, padx=6, pady=4)

        ex_left = ttk.Frame(ex_cols)
        ex_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        ex_right = ttk.Frame(ex_cols)
        ex_right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(16,0)) 

        EXCLUDE_COUNT = 5  # 总数量=5

        # 左列：排除1~3
        for i in range(3):
            rowf = ttk.Frame(ex_left)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"排除{i+1}").pack(side=tk.LEFT)
            ent = AttrPickerField(
                rowf, ATTRS, on_change=self._entry_changed, width=12,
                multi_mode=True,
                max_select=EXCLUDE_COUNT,
                get_pre=self._get_preselected_excludes,
                get_disabled=self._current_targets_set,
                apply_selected=self._apply_selected_excludes,
                picker_title="选择排除词条（多选）"
            )
            ent.pack(side=tk.LEFT, padx=(6,8))
            self.exclude_entries.append(ent)

        # 右列：排除4~5
        for i in range(3, EXCLUDE_COUNT):
            rowf = ttk.Frame(ex_right)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"排除{i+1}").pack(side=tk.LEFT)
            ent = AttrPickerField(
                rowf, ATTRS, on_change=self._entry_changed, width=12,
                multi_mode=True,
                max_select=EXCLUDE_COUNT,
                get_pre=self._get_preselected_excludes,
                get_disabled=self._current_targets_set,
                apply_selected=self._apply_selected_excludes,
                picker_title="选择排除词条（多选）"
            )
            ent.pack(side=tk.LEFT, padx=(6,8))
            self.exclude_entries.append(ent)


        ttk.Button(left, text="计算最佳组合", command=self.compute).pack(anchor="w", padx=6, pady=(0,10))

        # 左侧：已添加的模组
        listf = ttk.LabelFrame(left, text="已添加的模组")
        listf.pack(fill=tk.BOTH, expand=True)
        self.tree = ttk.Treeview(listf, columns=("desc",), show="headings", height=9)
        self.tree.heading("desc", text="模组词条 (0)")
        self.tree.column("desc", anchor=tk.CENTER, width=460, stretch=True)
        self.tree.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        btns = ttk.Frame(listf)
        btns.pack(fill=tk.X, padx=6, pady=(0,8))
        ttk.Button(btns, text="删除选中", command=self.delete_selected).pack(side=tk.LEFT)
        ttk.Button(btns, text="清空列表", command=self.clear_all).pack(side=tk.LEFT, padx=8)
        ttk.Button(btns, text="删除低级模组", command=self.delete_low_score_mods).pack(side=tk.LEFT, padx=8)

        # 右侧：计算结果
        resf = ttk.LabelFrame(right, text="计算结果")
        resf.pack(fill=tk.BOTH, expand=True)

        exportf = ttk.Frame(resf)
        exportf.pack(fill=tk.X, padx=8, pady=(8,8))
        ttk.Button(exportf, text="导出需求最佳（Word）",
                command=lambda: self.export_solutions_to_docx("req")).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(exportf, text="导出最高等级（Word）",
                command=lambda: self.export_solutions_to_docx("abs")).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(exportf, text="导出职业推荐（Word）",
                command=lambda: self.export_solutions_to_docx("bs")).pack(side=tk.LEFT, padx=(0,8))

        res_container = ttk.Frame(resf)
        res_container.pack(fill="both", expand=True)

        # 文本框
        req_sec, self.req_text = self._create_section(res_container, "目标需求方案(选择目标或排除后生效)")
        abs_sec, self.abs_text = self._create_section(res_container, "最高等级方案(对应战力最高)")
        bs_sec,  self.bs_text  = self._create_section(res_container, "职业属性方案(按照B站里的推荐属性)")

        # 用 grid 布局三等分
        res_container.rowconfigure(0, weight=1)
        res_container.rowconfigure(1, weight=1)
        res_container.rowconfigure(2, weight=1)
        res_container.columnconfigure(0, weight=1)

        abs_sec.grid(row=0, column=0, sticky="nsew")
        req_sec.grid(row=1, column=0, sticky="nsew")
        bs_sec.grid(row=2, column=0, sticky="nsew")

        self._refresh_mutex_states(False)

    def _create_section(self, parent, title: str):
        sec = ttk.Frame(parent)
        # 标题
        lbl = ttk.Label(sec, text=title, font=("Segoe UI", 10, "bold"))
        lbl.pack(anchor="w", padx=8, pady=(6, 2))
        # 文本 + 滚动条
        wrap = ttk.Frame(sec)
        wrap.pack(fill="both", expand=True, padx=8, pady=(0, 6))
        txt = tk.Text(wrap, wrap="word")
        sb  = ttk.Scrollbar(wrap, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        txt.configure(state="disabled")
        return sec, txt

    def _refresh_interfaces(self):
        """加载可用网卡到下拉框。"""
        try:
            ifaces = get_network_interfaces()  # [{'name','description','addresses',...}]:contentReference[oaicite:5]{index=5}
        except Exception as e:
            messagebox.showerror("错误", f"获取网卡失败：\n{e}")
            return

        items = []
        mapping = []
        for it in ifaces:
            name = it.get("name") or ""
            desc = it.get("description") or name
            addrs = ", ".join(a.get("addr","") for a in it.get("addresses", []))
            label = f"{desc}  ({addrs})" if addrs else desc
            items.append(label)
            mapping.append((label, name))

        self._iface_list = mapping
        self.cb_iface["values"] = items
        if items:
            self.cb_iface.current(0)
        else:
            self.var_iface.set("")

    def _start_sniff(self):
        """创建 PacketCapture 并启动抓包。"""
        if self.capture:
            messagebox.showwarning("提示", "抓包已在进行中。")
            return

        label = self.var_iface.get().strip()
        iface = ""
        for lab, real in self._iface_list:
            if lab == label:
                iface = real
                break

        # 未选则尝试 None（交给 scapy 自动选择）。你也可以强制要求必须选择。
        try:
            self.capture = PacketCapture(interface=iface or None)
        except Exception as e:
            messagebox.showerror("错误", f"创建抓包器失败：\n{e}")
            self.capture = None
            return

        # 启动，注册回调
        try:
            self._last_mod_time = time.time()
            self.capture.start_capture(callback=self._on_packet_data)  # 回调在子线程调用:contentReference[oaicite:6]{index=6}
        except Exception as e:
            messagebox.showerror("错误", f"启动抓包失败：\n{e}")
            self.capture = None
            return

        # UI 状态
        self.btn_stop.config(state="normal")
        messagebox.showinfo("提示", "已开始抓包。若长时间无数据，请确认已连接到游戏且网卡选择正确。")

        # 启动自动停止轮询
        self._schedule_auto_stop_check()

    def _stop_sniff(self):
        """停止抓包并清理。"""
        try:
            if self.capture:
                self.capture.stop_capture()
        finally:
            self.capture = None
            self.btn_stop.config(state="disabled")
            if self._auto_stop_job:
                try:
                    self.after_cancel(self._auto_stop_job)
                except Exception:
                    pass
                self._auto_stop_job = None
            messagebox.showinfo("提示", "抓包已停止。")

    def _schedule_auto_stop_check(self):
        """如果超过 _idle_timeout 秒没有新的模组，就自动停止。"""
        # 轮询自身
        def _tick():
            if not self.capture:
                return
            idle = time.time() - self._last_mod_time
            if idle >= self._idle_timeout:
                # 超时自动停
                try:
                    self.capture.stop_capture()
                finally:
                    self.capture = None
                    self.btn_stop.config(state="disabled")
                messagebox.showinfo("提示", f"{self._idle_timeout} 秒未收到新的模组，已自动停止。")
                return
            # 继续轮询
            self._auto_stop_job = self.after(1000, _tick)

        # 先取消旧的
        if self._auto_stop_job:
            try:
                self.after_cancel(self._auto_stop_job)
            except Exception:
                pass
        self._auto_stop_job = self.after(1000, _tick)

    def _on_packet_data(self, data: dict):
        """
        抓包回调（子线程）。data 形如 {'v_data': <CharSerialize.VData>}:contentReference[oaicite:7]{index=7}
        这里把 VData 交给 ModuleParser 解析出模组，再把结果合并进 UI 的 mods 列表。
        """
        try:
            v_data = data.get('v_data')
            if v_data is None:
                return
            # 解析模组列表
            modules = self.module_parser.parse_module_info(v_data, category="攻击", attributes=None) or []
        except Exception:
            return

        # 转换为 mod2 期待的结构：{'a1','p1','a2','p2', 可选 'a3','p3'}
        # ModulePart.value 是这条词条的“点数/链接数”，这里直接当作点数使用
        new_mods = []
        for m in modules:
            if not getattr(m, "parts", None):
                continue
            # 取前 2~3 条
            ps = list(m.parts)
            if len(ps) >= 2:
                g = {"a1": ps[0].name, "p1": int(ps[0].value),
                     "a2": ps[1].name, "p2": int(ps[1].value)}
                if len(ps) >= 3:
                    g["a3"] = ps[2].name
                    g["p3"] = int(ps[2].value)
                new_mods.append(g)

        if not new_mods:
            return

        # 更新“最近收到时间”，用于自动停止
        self._last_mod_time = time.time()

        self.after(0, lambda: self._append_sniffed_mods(new_mods))

    def _append_sniffed_mods(self, mods_from_net):
        """把抓到的新模组加入列表与 Treeview。"""
        added = 0
        for g in mods_from_net:
            # 只做最基本校验：词条名在候选列表里
            a1, a2 = g.get("a1",""), g.get("a2","")
            if a1 in ATTRS and a2 in ATTRS:
                self.mods.append(g)
                self.tree.insert("", tk.END, values=(format_mod_str(g),))
                added += 1
        if added:
            try:
                self.title(f"模组词条配点计算器   by椛椿   ｜ 实时导入 +{added}")
            except Exception:
                pass
        self._update_mod_counters()

    def delete_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        idxs = sorted([self.tree.index(i) for i in sel], reverse=True)
        for i in idxs:
            del self.mods[i]
        for i in sel:
            self.tree.delete(i)
        self._update_mod_counters()

    def clear_all(self):
        if messagebox.askyesno("确认", "确定清空所有模组？"):
            self.mods.clear()
            for i in self.tree.get_children():
                self.tree.delete(i)
            self._update_mod_counters()

    # 计算（异步 + 进度条，不截断）
    def compute(self):
        # 固定为 4 槽
        slots = 4

        # 基本校验
        n = len(self.mods)
        if n < slots:
            messagebox.showerror("错误", f"需要至少 {slots} 颗模组。当前：{n}")
            return

        # 读取职业词条集合
        job = self.job_field.get_value() if hasattr(self, "job_field") else ""
        prof_attrs = CLASS_ATTRS.get(job, set()) if job and job != "（不选）" else set()


        # 读取目标（校验）
        targets, target_set = [], set()
        for ent, sp in self.targets:
            name = ent.get_value()
            try:
                req = int(sp.get())
            except ValueError:
                req = 0
            if name and name not in ATTRS:
                messagebox.showerror("错误", f"目标词条“{name}”不在列表中。")
                return
            if req < 0 or req > 6:
                messagebox.showerror("错误", "期望等级范围应为 0~6（0 表示不指定）。")
                return
            if name:
                targets.append((name, req))
                target_set.add(name)

        # 读取排除（校验）
        exclude_set = set()
        if hasattr(self, "exclude_entries"):
            for ent in self.exclude_entries:
                name = ent.get_value()
                if name:
                    if name not in ATTRS:
                        messagebox.showerror("错误", f"排除词条“{name}”不在列表中。")
                        return
                    exclude_set.add(name)

        # 互斥校验：目标 vs 排除
        inter = target_set & exclude_set
        if inter:
            messagebox.showerror("冲突", f"目标与排除不能相同：{', '.join(sorted(inter))}")
            return
        
        # 当目标与排除都为空时，抑制“需求最佳方案”
        self._req_suppress = (len(targets) == 0 and len(exclude_set) == 0)

        # 过滤掉含排除词条的模组，缩小搜索空间
        filtered = []
        for g in self.mods:
            if exclude_set and (g["a1"] in exclude_set or g["a2"] in exclude_set):
                continue
            filtered.append(g)

        m = len(filtered)
        if m < slots:
            messagebox.showerror("错误", f"需要至少 {slots} 颗模组（排除后不足）。当前：{m}")
            return

        # 估算组合总数（用于进度）
        try:
            import math
            total_combos = math.comb(m, slots)
        except Exception:
            def _comb(n, k):
                if k < 0 or k > n: return 0
                if k == 0 or k == n: return 1
                k = min(k, n-k)
                num = 1
                den = 1
                for i in range(1, k+1):
                    num *= n - (k - i)
                    den *= i
                return num // den
            total_combos = _comb(m, slots)

        # 进度更新步长（约 200 次进度更新）
        COMB_UPDATE_EVERY = max(2000, total_combos // 200)

        # 若已有任务在跑，阻止并行
        if getattr(self, "_worker", None) and self._worker.is_alive():
            messagebox.showwarning("提示", "已有计算正在进行。请先取消或等待完成。")
            return

        # 记录当前池以避免展示索引错位 & 启动后台线程 + 进度窗
        import threading
        self._cancel = False
        self._pool_current = filtered

        dlg = _Progress(self, total_combos, on_cancel=lambda: self._set_cancel(dlg))

        self._worker = threading.Thread(
            target=self._compute_mp,
            args=(filtered, slots, targets, total_combos, prof_attrs),
            daemon=True
        )
        self._worker.start()

        # 轮询线程结果
        self._poll_compute(dlg)

    def _set_cancel(self, dlg):
        self._cancel = True
        try:
            dlg.var.set("正在取消…")
        except Exception:
            pass

    def _poll_compute(self, dlg):
        try:
            while True:
                typ, payload = self._q.get_nowait()
                if typ == "progress":
                    done, total = payload
                    try:
                        dlg.set_progress(done, total)
                    except Exception:
                        pass
                elif typ == "done":
                    try:
                        dlg.destroy()
                    except Exception:
                        pass
                    req_list, abs_list, bs_list, max_c6, scanned, total = payload
                    self._render_results(req_list, abs_list, bs_list, max_c6, scanned, total)
                elif typ == "cancelled":
                    try:
                        dlg.destroy()
                    except Exception:
                        pass
                    messagebox.showinfo("已取消", "计算已取消。")
        except queue.Empty:
            pass

        if self._worker and self._worker.is_alive():
            self.after(80, lambda: self._poll_compute(dlg))

    def _compute_worker(self, mods, slots, targets, total, update_every, prof_attrs):
        processed = 0
        import heapq
        req_top = []  # (rank_key, pack)
        bs_top  = []  # (battle_key, pack)
        abs_top = []  # (abs_key, pack)  # 绝对榜键：(-c6, -sumlv, waste, -total_points)

        def keep_top(heap, item, limit):
            if len(heap) < limit:
                heapq.heappush(heap, item)
            else:
                # 更优（更小）则替换
                if item[0] < heap[0][0]:
                    heapq.heapreplace(heap, item)

        try:
            import itertools
            summarize_levels_local = summarize_levels
            rank_tuple_local = rank_tuple

            m = len(mods)
            it = itertools.combinations(range(m), slots)

            for idxs in it:
                if self._cancel:
                    self._q.put(("cancelled", None))
                    return

                # 兼容第三词条的聚合点数
                pts = {}
                for i in idxs:
                    g = mods[i]
                    pts[g["a1"]] = pts.get(g["a1"], 0) + g["p1"]
                    pts[g["a2"]] = pts.get(g["a2"], 0) + g["p2"]
                    if "a3" in g and "p3" in g:
                        pts[g["a3"]] = pts.get(g["a3"], 0) + g["p3"]

                levels, c6, sumlv, waste = summarize_levels_local(pts)
                total_points = sum(pts.values())

                # 统一打包（包含实际 mod 对象，避免索引错位）
                combo_idx = tuple(idxs)
                combo_mods = tuple(mods[i] for i in idxs)
                pack = (None, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points)

                # 1) 需求最佳：用 rank_tuple（含目标惩罚/换算分/等级&点数总和）
                req_key = rank_tuple_local(levels, c6, sumlv, waste, targets, pts_map=pts)
                keep_top(req_top, (req_key, pack), REQ_KEEP_LIMIT)

                # 2) 绝对榜：6多→总等级高→浪费少→点数总和高
                abs_key = (-c6, -sumlv, waste, -total_points)
                keep_top(abs_top, (abs_key, pack), ABS_KEEP_LIMIT)

                # 3) 战力分：6多→职业词条命中多→再用 req_key
                bs_key = _battle_key(levels, req_key, prof_attrs)
                keep_top(bs_top, (bs_key, pack), BS_KEEP_LIMIT)

                processed += 1
                if processed % update_every == 0 or processed == total:
                    self._q.put(("progress", (processed, total)))

            # 小顶堆转升序列表（已是从优到劣）
            req_list = [x[1] for x in sorted(req_top, key=lambda t: t[0])]
            abs_list = [x[1] for x in sorted(abs_top, key=lambda t: t[0])]
            bs_list  = [x[1] for x in sorted(bs_top,  key=lambda t: t[0])]

            max_c6 = max((p[5] for p in abs_list), default=0)
            self._q.put(("done", (req_list, abs_list, bs_list, max_c6, processed, total)))

        except Exception:
            self._q.put(("done", ([], [], [], -1, processed, total)))

    def _render_results(self, req_list, abs_list, bs_list, max_c6, scanned, total_combos):
        # 保存给导出用
        self._last_req_list = list(req_list)
        self._last_abs_list = list(abs_list)
        self._last_bs_list  = list(bs_list)

        def lv6_attrs(levels):
            arr = [a for a, lv in levels.items() if lv == 6]
            return "，".join(arr) if arr else "（无）"

        def combo_lines(combo_mods):
            return "\n".join([format_mod_str(g) for g in combo_mods])

        # ========== 需求最佳 ==========
        # 判断是否需要抑制：目标与排斥都为空 → 直接清空该面板
        has_target = any(ent.get_value() for ent, _ in self.targets)
        has_exclude = any(ent.get_value() for ent in self.exclude_entries)
        suppress_req = (not has_target) and (not has_exclude)

        self.req_text.configure(state="normal")
        self.req_text.delete("1.0", tk.END)

        if suppress_req:
            # 完全置空；并清空导出缓存，避免导出产生内容
            self._last_req_list = []
            self.req_text.configure(state="disabled")
        else:
            if not req_list:
                self.req_text.insert(tk.END, "（无可行组合，请检查槽位、排除条件与已添加模组）")
            else:
                # 根据是否有目标动态说明
                if has_target:
                    desc = "（按目标约束排序：目标更接近→换算分高→6更多→≥5向量更优→等级总和高→点数总和高）"
                else:
                    desc = "（无目标：换算分高→6更多→≥5向量更优→等级总和高→点数总和高）"
                self.req_text.insert(tk.END, desc + "\n\n")
                for k, pack in enumerate(req_list[:REQ_KEEP_LIMIT], start=1):
                    _, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points = pack
                    self.req_text.insert(tk.END, f"方案 #{k}\n")
                    self.req_text.insert(tk.END, f"{combo_lines(combo_mods)}\n")
                    self.req_text.insert(tk.END, f"等级概览：\n{levels_pretty(levels)}\n")
                    self.req_text.insert(tk.END, f"6级词条：{lv6_attrs(levels)}\n")
                    self.req_text.insert(tk.END, f"统计：6级数量={c6}，等级总和={sumlv}，浪费点={waste}，点数总和={total_points}\n\n")

            # 只有在不抑制时，才追加进度尾注
            if not suppress_req:
                self.req_text.insert(tk.END, f"（已扫描组合：{scanned:,} / 总计 {total_combos:,}）")
            self.req_text.configure(state="disabled")

        # 最高等级（Top 100）
        self.abs_text.configure(state="normal")
        self.abs_text.delete("1.0", tk.END)
        if not abs_list:
            self.abs_text.insert(tk.END, "（无 6 级可达，已在需求方案中给出最优）")
        else:
            self.abs_text.insert(tk.END, f"（绝对最高 6级数量 = {max_c6}；共 {len(abs_list)} 个候选，显示前 {min(len(abs_list), ABS_KEEP_LIMIT)} 个）\n\n")
            for k, pack in enumerate(abs_list[:ABS_KEEP_LIMIT], start=1):
                _, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points = pack
                self.abs_text.insert(tk.END, f"方案 #{k}\n")
                self.abs_text.insert(tk.END, f"{combo_lines(combo_mods)}\n")
                self.abs_text.insert(tk.END, f"等级概览：\n{levels_pretty(levels)}\n")
                self.abs_text.insert(tk.END, f"6级词条：{lv6_attrs(levels)}\n")
                self.abs_text.insert(tk.END, f"统计：6级数量={c6}，等级总和={sumlv}，浪费点={waste}，点数总和={total_points}\n\n")
        self.abs_text.configure(state="disabled")

        # ========== 战力分推荐（Top 100） ==========
        self.bs_text.configure(state="normal")
        self.bs_text.delete("1.0", tk.END)

        # 单选职业获取
        job = self.job_field.get_value() if hasattr(self, "job_field") else ""
        job_selected = bool(job and job != "无" and job != "（不选）")


        if not job_selected:
            # 未选择职业或选择了“无”→不展示战力分推荐
            self._last_bs_list = []
            self.bs_text.configure(state="disabled")
        else:
            if not bs_list:
                self.bs_text.insert(tk.END, "（无可行组合）")
            else:
                head = f"（职业：{job}；排序：6级更多→职业词条命中更多→需求排序）\n\n"
                self.bs_text.insert(tk.END, head)

                prof_attrs = CLASS_ATTRS.get(job, set())
                for k, pack in enumerate(bs_list[:BS_KEEP_LIMIT], start=1):
                    _, combo_idx, combo_mods, pts, levels, c6, sumlv, waste, total_points = pack
                    self.bs_text.insert(tk.END, f"方案 #{k}\n")
                    self.bs_text.insert(tk.END, f"{combo_lines(combo_mods)}\n")
                    self.bs_text.insert(tk.END, f"等级概览：\n{levels_pretty(levels)}\n")
                    self.bs_text.insert(tk.END, f"6级词条：{lv6_attrs(levels)}\n")
                    match = sum(1 for a, lv in levels.items() if lv > 0 and a in prof_attrs) if prof_attrs else 0
                    self.bs_text.insert(tk.END, f"统计：6级数量={c6}，职业命中={match}，等级总和={sumlv}，浪费点={waste}，点数总和={total_points}\n\n")
            self.bs_text.configure(state="disabled")


if __name__ == "__main__":
    import multiprocessing as mp
    mp.freeze_support()
    try:
        mp.set_start_method("spawn")
    except RuntimeError:
        pass

    import sys
    if getattr(sys, "frozen", False):
        mp.set_executable(sys.executable)

    App().mainloop()
