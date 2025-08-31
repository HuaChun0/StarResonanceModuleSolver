import sys
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
from tkinter.scrolledtext import ScrolledText
import threading, queue
from module_compute import ModuleOptimizerCalculations
from typing import Dict, Any
import openpyxl
import time
from typing import Dict, Any, List, Tuple

#抓包模块
from module_parser import ModuleParser
from packet_capture import PacketCapture
from network_interface_util import get_network_interfaces, select_network_interface

def levels_pretty(levels: dict) -> str:
    items = [(a, lv) for a, lv in levels.items() if lv > 0]
    return "，".join([f"{a}: Lv{lv}" for a, lv in items]) if items else "（无有效等级）"

try:
    from ctypes import windll
    windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

BASE_DIR = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))

try:
    import pytesseract  # type: ignore
    HAS_TESS = True
except Exception:
    pytesseract = None
    HAS_TESS = False

TESS_DIR = BASE_DIR / "tesseract"
TESS_EXE = TESS_DIR / ("tesseract.exe" if os.name == "nt" else "tesseract")
TESSDATA_DIR = TESS_DIR / "tessdata"

if HAS_TESS and TESS_EXE.exists() and TESSDATA_DIR.exists():
    pytesseract.pytesseract.tesseract_cmd = str(TESS_EXE)
    os.environ.pop("TESSDATA_PREFIX", None)
    os.environ["TESSDATA_PREFIX"] = str(TESSDATA_DIR)
else:
    HAS_TESS = False

# ==== imports ====
import sys
from pathlib import Path
try:
    import pandas as pd  # 可选
    HAS_PANDAS = True
except Exception:
    HAS_PANDAS = False

# 新版 vision_recog
import vision_recog as vr
from vision_recog import (
    load_icon_templates,
    parse_screenshot_to_gems,
    parse_folder_to_gems,
    configure_tesseract,
    set_ocr_geometry,
    ocr_is_available,
)

# ==== 路径与常量 ====
BASE_DIR = Path(__file__).resolve().parent
ICON_DIR = BASE_DIR / "icons"

configure_tesseract(TESS_EXE, TESSDATA_DIR)
print("OCR available:", ocr_is_available())

# 2) 适配几何（右侧 ROI 的起始偏移与宽度；与 mod1 一样的默认值：0.12, 1.40）
set_ocr_geometry(pad_ratio=0.12, right_w_factor=1.40)

vr.USE_BAR_CROP = True
vr.USE_LEFTMOST_TOKEN = True

# 可选依赖：pandas
try:
    import pandas as pd  # noqa: F401
    HAS_PANDAS = True
except Exception:
    HAS_PANDAS = False


#== 常量==
ATTRS = [
    "抵御魔法","抵御物理","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取",
    "力量加持","敏捷加持","智力加持","特攻伤害加持","精英打击","特攻治疗加持",
    "专精治疗加持","施法专注","攻速专注","暴击专注","幸运专注","极·伤害叠加",
    "极·灵活身法","极·急救措施","极·全队幸暴"
]
POINT_MIN, POINT_MAX = 1, 10
SLOT_DEFAULT = 4
REQ_KEEP_LIMIT = 100   # 需求最佳
BS_KEEP_LIMIT  = 100   # 战力分推荐
ABS_KEEP_LIMIT = 100   # 绝对最高方案

# 职业对应词条
CLASS_ATTRS = {
    "光盾": {"抵御魔法","抵御物理","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取"},
    "防回": {"抵御魔法","抵御物理","暴击专注","极·生命凝聚","极·绝境守护","极·生命波动","极·生命吸取"},
    "岩盾": {"抵御魔法","抵御物理","极·绝境守护"},
    "格挡": {"抵御魔法","抵御物理","极·绝境守护","幸运专注"},
    "惩击": {"智力加持","特攻治疗加持","专精治疗加持","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "愈合": {"施法专注","特攻治疗加持","专精治疗加持","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "狂音": {"智力加持","特攻治疗加持","专精治疗加持","攻速专注","幸运专注","极·生命凝聚","极·急救措施","极·全队幸暴"},
    "协奏": {"特攻治疗加持","专精治疗加持","施法专注","极·生命凝聚","极·急救措施","极·全队幸暴","暴击专注"},
    "居合": {"敏捷加持","特攻伤害加持","精英打击","暴击专注","极·伤害叠加","极·灵活身法","极·生命波动"},
    "月刃": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","幸运专注","极·伤害叠加","极·灵活身法","极·生命波动"},
    "重装": {"特攻伤害加持","精英打击","攻速专注","力量加持","极·伤害叠加","极·灵活身法","极·生命波动"},
    "空战": {"特攻伤害加持","精英打击","力量加持","极·伤害叠加","极·灵活身法","极·生命波动","暴击专注","幸运专注"},
    "冰矛": {"智力加持","特攻伤害加持","精英打击","施法专注","暴击专注","幸运专注","极·伤害叠加","极·灵活身法"},
    "射线": {"智力加持","特攻伤害加持","精英打击","施法专注","极·伤害叠加","极·灵活身法"},
    "驭兽": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","极·伤害叠加","极·灵活身法"},
    "驯鹰": {"敏捷加持","特攻伤害加持","精英打击","攻速专注","暴击专注","幸运专注","极·伤害叠加","极·灵活身法"},
}

CLASS_LIST = list(CLASS_ATTRS.keys())

#== 等级阈值/评分==
def points_to_level(p: int) -> int:
    if p <= 0: return 0
    if p <= 3: return 1
    if p <= 7: return 2
    if p <= 11: return 3
    if p <= 15: return 4
    if p <= 19: return 5
    return 6  # >=20

def format_gem_str(g):
    parts = [f"{g['a1']} {g['p1']}点", f"{g['a2']} {g['p2']}点"]
    if "a3" in g and "p3" in g:
        parts.append(f"{g['a3']} {g['p3']}点")
    return "  +  ".join(parts)


#== 等级展示辅助==
def _levels_from_points(pts: dict) -> dict:
    """将属性累计点数映射为等级（沿用项目里的 points_to_level 规则）。"""
    return {a: points_to_level(v) for a, v in pts.items()}

def _pretty_levels_table(levels: dict) -> str:
    """把 {属性: 等级} 渲染成多行文本，只显示 >=1 级的属性。"""
    if not levels:
        return "（无）"
    rows = [(a, lv) for a, lv in levels.items() if lv and lv > 0]
    rows.sort(key=lambda x: (-x[1], x[0]))  # 先按等级降序，再按名称升序
    if not rows:
        return "（全为0级）"
    return "\n".join([f" - {a}：{lv}级" for a, lv in rows])


ICON_DIR = BASE_DIR / "icons"

from vision_recog import load_icon_templates

def _load_icon_templates():
    return load_icon_templates(ICON_DIR)


class NetworkInterfaceDialog(tk.Toplevel):
    """网络接口选择对话框"""
    def __init__(self, master, interfaces):
        super().__init__(master)
        self.title("选择网络接口")
        self.resizable(False, False)
        self.grab_set()
        self.result = None
        self.interfaces = interfaces
        
        # 设置窗口大小和位置
        self.geometry("600x400")
        self.update_idletasks()
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except:
            pass
        
        self.build_ui()
        
    def build_ui(self):
        # 主框架
        main_frame = ttk.Frame(self, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        ttk.Label(main_frame, text="请选择要用于抓包的网络接口：", 
                 font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(0,10))
        
        # 接口列表框架
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0,10))
        
        # 列表框和滚动条
        self.listbox = tk.Listbox(list_frame, height=12, font=("Segoe UI", 9))
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.listbox.yview)
        self.listbox.configure(yscrollcommand=scrollbar.set)
        
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 填充接口信息
        self.populate_interfaces()
        
        # 按钮框架
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(10,0))
        
        ttk.Button(btn_frame, text="确定", command=self.on_ok).pack(side=tk.RIGHT, padx=(5,0))
        ttk.Button(btn_frame, text="取消", command=self.on_cancel).pack(side=tk.RIGHT)
        
        # 双击选择
        self.listbox.bind("<Double-Button-1>", lambda e: self.on_ok())
        
        # ESC键取消
        self.bind("<Escape>", lambda e: self.on_cancel())
        self.protocol("WM_DELETE_WINDOW", self.on_cancel)
        
    def populate_interfaces(self):
        """填充网络接口列表"""
        from network_interface_util import find_default_network_interface
        
        default_index = find_default_network_interface(self.interfaces)
        
        for i, interface in enumerate(self.interfaces):
            name = interface['name']
            description = interface.get('description', name)
            is_up = "✓" if interface.get('is_up', False) else "✗"
            addresses = [addr['addr'] for addr in interface['addresses']]
            addr_str = ", ".join(addresses) if addresses else "无IP地址"
            
            # 格式化显示文本
            status_text = "活动" if interface.get('is_up', False) else "非活动"
            default_text = " (默认)" if i == default_index else ""
            
            display_text = f"{i:2d}. {is_up} {description}{default_text}"
            detail_text = f"     地址: {addr_str} | 状态: {status_text}"
            
            # 添加到列表框
            self.listbox.insert(tk.END, display_text)
            self.listbox.insert(tk.END, detail_text)
            self.listbox.insert(tk.END, "")  # 空行分隔
            
        # 默认选择第一个接口（如果有的话）
        if self.interfaces:
            select_index = default_index if default_index is not None else 0
            # 计算在listbox中的实际行号（每个接口占3行：名称、详情、空行）
            listbox_index = select_index * 3
            self.listbox.selection_set(listbox_index)
            self.listbox.see(listbox_index)
            
    def on_ok(self):
        """确定选择"""
        selection = self.listbox.curselection()
        if not selection:
            tk.messagebox.showwarning("提示", "请选择一个网络接口！")
            return
            
        # 计算实际的接口索引（每个接口占3行）
        selected_line = selection[0]
        interface_index = selected_line // 3
        
        if 0 <= interface_index < len(self.interfaces):
            self.result = interface_index
            self.destroy()
        else:
            tk.messagebox.showerror("错误", "选择的接口无效！")
            
    def on_cancel(self):
        """取消选择"""
        self.result = None
        self.destroy()

class _Progress(tk.Toplevel):
    def __init__(self, master, total, on_cancel):
        super().__init__(master)
        self.title("正在计算…")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", on_cancel)
        self.grab_set()  # 模态
        self.var = tk.StringVar(value="准备中…")
        ttk.Label(self, textvariable=self.var).pack(padx=14, pady=(12,6))
        self.pb = ttk.Progressbar(self, length=360, mode="determinate", maximum=max(1, total))
        self.pb.pack(padx=14, pady=6)
        ttk.Button(self, text="取消", command=on_cancel).pack(pady=(6,12))
        self.update_idletasks()
        # 居中
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def set_progress(self, done, total):
        self.pb["value"] = min(done, total)
        self.var.set(f"已处理 {done:,} / {total:,} 个组合…")

#== 主应用==
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("模组求解器器   by椛椿")
        self.geometry("1650x1000")
        self.minsize(1080, 680)
        self.gems = []  # 每项: {"a1":..., "p1":int, "a2":..., "p2":int}

        # 记录互斥输入框引用
        self.targets = []
        self.target_entries = []
        self.exclude_entries = []

        # 异步计算通信（仅用于 OCR 文件夹识别；旧 compute worker 已移除）
        self._cancel = False

        # 添加抓包相关变量
        self.packet_capture = None
        self.is_capturing = False
        self.auto_pause_after_data = True
        self.auto_pause_timer = None
        self.last_data_time = None

        self.build_ui()
        self._dedup_set = set()

    def _set_cancel(self, dlg):
        """设置取消标志"""
        self._cancel = True
        try:
            if dlg and dlg.winfo_exists():
                dlg.destroy()
        except:
            pass

    def toggle_packet_capture(self):
        """切换抓包状态"""
        if not self.is_capturing:
            self.start_packet_capture()
        else:
            self.stop_packet_capture()

    def start_packet_capture(self):
        """开始抓包 - 带网卡选择功能"""
        try:
            # 获取网络接口
            interfaces = get_network_interfaces()
            if not interfaces:
                messagebox.showerror("错误", "未找到可用的网络接口")
                return
            
            # 如果只有一个接口，直接使用
            if len(interfaces) == 1:
                selected_interface = interfaces[0]['name']
                messagebox.showinfo("网络接口", f"检测到唯一网络接口，将使用: {selected_interface}")
            else:
                # 显示接口选择对话框
                dialog = NetworkInterfaceDialog(self, interfaces)
                self.wait_window(dialog)
                
                if dialog.result is None:
                    # 用户取消选择
                    return
                    
                selected_interface = interfaces[dialog.result]['name']
            
            # 创建并启动抓包器
            self.packet_capture = PacketCapture(interface=selected_interface)
            self.packet_capture.start_capture(callback=self.on_packet_data_received)
            
            self.is_capturing = True
            self.capture_btn.config(text="停止抓包")
            self.capture_status.config(text=f"抓包中 - {selected_interface}")
            
            messagebox.showinfo("开始抓包", f"已开始监听网络接口: {selected_interface}\n请在游戏中切换地图或重新登陆")
            
        except Exception as e:
            messagebox.showerror("抓包错误", f"启动抓包失败：\n{e}")

    def stop_packet_capture(self):
        """停止抓包"""
        if self.packet_capture:
            self.packet_capture.stop_capture()
            self.packet_capture = None
        
        self.is_capturing = False
        self.capture_btn.config(text="开始抓包")
        self.capture_status.config(text="未开始")

    def on_packet_data_received(self, data: Dict[str, Any]):
        """处理接收到的抓包数据"""
        try:
            if 'v_data' not in data:
                return
            
            v_data = data['v_data']
            
            # 记录数据获取时间
            import time
            self.last_data_time = time.time()
            
            # 解析模组数据
            parser = ModuleParser()
            
            # 调用解析方法但不让它自动优化和退出
            modules = self._parse_modules_from_vdata(v_data)
            
            if modules:
                # 转换为UI格式并添加到gems列表
                converted_count = self._convert_modules_to_ui_format_smart_dedup(modules)
                
                if converted_count > 0:
                    # 更新状态显示
                    self.capture_status.config(text=f"已获取 {converted_count} 个模组")
                    messagebox.showinfo("抓包成功", f"成功解析 {converted_count} 个模组数据\n已自动添加到模组列表")
                    
                    # 如果启用了自动暂停，则暂停抓包
                    if self.auto_pause_after_data:
                        self.after(2000, self._auto_pause_capture)  # 2秒后自动暂停
                
        except Exception as e:
            print(f"处理抓包数据时出错：{e}")

    def _auto_pause_capture(self):
        """自动暂停抓包"""
        if self.is_capturing:
            self.stop_packet_capture()
            self.capture_status.config(text="已自动暂停 - 数据获取完成")
            
            # 可选：显示提示信息
            messagebox.showinfo("自动暂停", "检测到模组数据已获取完成，抓包已自动暂停。\n如需继续抓包请重新点击开始抓包。")

    def _parse_modules_from_vdata(self, v_data):
        """从v_data解析模组，不触发优化和退出"""
        try:
            mod_infos = v_data.Mod.ModInfos
            modules = []
            
            for package_type, package in v_data.ItemPackage.Packages.items():
                for key, item in package.Items.items():
                    if item.HasField('ModNewAttr') and item.ModNewAttr.ModParts:
                        config_id = item.ConfigId
                        
                        # 导入模组名称映射
                        try:
                            from module_types import MODULE_NAMES, MODULE_ATTR_NAMES
                        except ImportError:
                            MODULE_NAMES = {}
                            MODULE_ATTR_NAMES = {}
                        
                        module_name = MODULE_NAMES.get(config_id, f"未知模组({config_id})")
                        mod_parts = list(item.ModNewAttr.ModParts)
                        
                        # 查找模组详细信息
                        mod_info = mod_infos.get(key) if mod_infos else None
                        
                        if mod_info and hasattr(mod_info, 'InitLinkNums'):
                            from module_types import ModuleInfo, ModulePart
                            
                            module_info = ModuleInfo(
                                name=module_name,
                                config_id=config_id,
                                uuid=item.Uuid,
                                quality=item.Quality,
                                parts=[]
                            )

                            init_link_nums = mod_info.InitLinkNums
                            for i, part_id in enumerate(mod_parts):
                                if i < len(init_link_nums):
                                    attr_name = MODULE_ATTR_NAMES.get(part_id, f"未知属性({part_id})")
                                    attr_value = init_link_nums[i]
                                    module_part = ModulePart(
                                        id=part_id,
                                        name=attr_name,
                                        value=attr_value
                                    )
                                    module_info.parts.append(module_part)
                            modules.append(module_info)
                    
            return modules
            
        except Exception as e:
            print(f"解析模组数据失败：{e}")
            return []

    def _convert_modules_to_ui_format(self, modules):
        """将ModuleInfo对象转换为UI字典格式并添加到gems列表"""
        converted_count = 0
        
        for module in modules:
            if len(module.parts) < 2:
                continue  # 跳过少于2个词条的模组
            
            # 转换为UI字典格式
            gem_dict = {}
            
            # 处理前两个词条（必需）
            if len(module.parts) >= 1:
                gem_dict["a1"] = module.parts[0].name
                gem_dict["p1"] = module.parts[0].value
            
            if len(module.parts) >= 2:
                gem_dict["a2"] = module.parts[1].name
                gem_dict["p2"] = module.parts[1].value
            
            # 处理第三个词条（可选）
            if len(module.parts) >= 3:
                gem_dict["a3"] = module.parts[2].name
                gem_dict["p3"] = module.parts[2].value
            
            # 验证属性名是否在已知列表中
            valid_attrs = 0
            for attr_key in ["a1", "a2", "a3"]:
                if attr_key in gem_dict and gem_dict[attr_key] in ATTRS:
                    valid_attrs += 1
            
            # 至少需要前两个属性有效
            if valid_attrs >= 2 and gem_dict.get("a1") in ATTRS and gem_dict.get("a2") in ATTRS:
                # 检查是否重复（基于属性组合）
                gem_signature = self._get_gem_signature(gem_dict)
                if gem_signature not in self._dedup_set:
                    self._dedup_set.add(gem_signature)
                    self.gems.append(gem_dict)
                    self.tree.insert("", tk.END, values=(format_gem_str(gem_dict),))
                    converted_count += 1
        
        return converted_count

    def _get_gem_signature(self, gem_dict):
        """生成模组签名用于去重"""
        attrs = []
        if "a1" in gem_dict and "p1" in gem_dict:
            attrs.append(f"{gem_dict['a1']}:{gem_dict['p1']}")
        if "a2" in gem_dict and "p2" in gem_dict:
            attrs.append(f"{gem_dict['a2']}:{gem_dict['p2']}")
        if "a3" in gem_dict and "p3" in gem_dict:
            attrs.append(f"{gem_dict['a3']}:{gem_dict['p3']}")
        return "|".join(sorted(attrs))

    def _convert_modules_to_ui_format_smart_dedup(self, modules):
        """智能去重：相同配置最多保留3个"""
        converted_count = 0
        attr_combo_count = {}  # 记录每种属性组合的数量
        
        for module in modules:
            if len(module.parts) < 2:
                continue  # 跳过少于2个词条的模组
            
            # 转换为UI字典格式
            gem_dict = {}
            
            # 处理词条
            if len(module.parts) >= 1:
                gem_dict["a1"] = module.parts[0].name
                gem_dict["p1"] = module.parts[0].value
            if len(module.parts) >= 2:
                gem_dict["a2"] = module.parts[1].name
                gem_dict["p2"] = module.parts[1].value
            if len(module.parts) >= 3:
                gem_dict["a3"] = module.parts[2].name
                gem_dict["p3"] = module.parts[2].value
            
            # 验证属性名
            if (gem_dict.get("a1") in ATTRS and gem_dict.get("a2") in ATTRS and 
                (not gem_dict.get("a3") or gem_dict.get("a3") in ATTRS)):
                
                # 生成属性组合签名
                signature = self._get_gem_signature(gem_dict)
                current_count = attr_combo_count.get(signature, 0)
                
                # 相同配置最多保留3个
                if current_count < 3:
                    attr_combo_count[signature] = current_count + 1
                    self.gems.append(gem_dict)
                    self.tree.insert("", tk.END, values=(format_gem_str(gem_dict),))
                    converted_count += 1
        
        return converted_count

    def _open_group_picker_addgem(self):
        """添加模组：整组多选，最多3个；至少2个在 add_gem() 校验"""
        current = [x for x in [self.a1.get_value(), self.a2.get_value(), self.a3.get_value()] if x]
        dlg = MultiCheckDialog(self, "选择模组词条（最多3个）", ATTRS, preselected=current, max_select=3, columns=4)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        # 回填到 a1/a2/a3
        sel = dlg.result[:3]
        # 先清空
        for f in (self.a1, self.a2, self.a3):
            f.var.set("")
        # 依序填入
        if len(sel) > 0: self.a1.var.set(sel[0])
        if len(sel) > 1: self.a2.var.set(sel[1])
        if len(sel) > 2: self.a3.var.set(sel[2])
        self._refresh_mutex_states(False)

    def _open_group_picker_targets(self):
        """期望属性（6个）：整组多选"""
        current = [ent.get_value() for ent, _ in self.targets if ent.get_value()]
        dlg = MultiCheckDialog(self, "选择期望属性", ATTRS, preselected=current, max_select=6, columns=4)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        sel = dlg.result[:6]
        # 依序填入，不够的清空
        for i, (ent, _) in enumerate(self.targets):
            ent.var.set(sel[i] if i < len(sel) else "")
        self._refresh_mutex_states(True)

    def _open_group_picker_excludes(self):
        """排除属性（9个）：整组多选"""
        current = [ent.get_value() for ent in self.exclude_entries if ent.get_value()]
        dlg = MultiCheckDialog(self, "选择排除属性", ATTRS, preselected=current, max_select=9, columns=4)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        sel = dlg.result[:9]
        for i, ent in enumerate(self.exclude_entries):
            ent.var.set(sel[i] if i < len(sel) else "")
        self._refresh_mutex_states(True)

    def _open_group_picker_job(self):
        """职业：单选（最多1个）"""
        cur = self.job_field.get_value()
        pre = [cur] if cur else []
        dlg = MultiCheckDialog(self, "选择职业", CLASS_LIST, preselected=pre, max_select=1, columns=4)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self.job_field.var.set(dlg.result[0] if dlg.result else "")


    def export_solutions_to_docx(self, which: str) -> None:
        """
        导出方案为 Word（.docx）。
        which ∈ {"req", "abs", "bs"} 分别对应：需求最佳 / 最高等级 / 战力分推荐。
        依赖：python-docx  (pip install python-docx)
        """
        # 1) 依赖检查
        try:
            from docx import Document
        except Exception:
            from tkinter import messagebox
            messagebox.showerror("缺少依赖", "需要安装 python-docx 才能导出 Word：\n\npip install python-docx")
            return

        # 2) 选择数据源
        if which == "req":
            data = getattr(self, "_last_req_list", [])
            title = "需求最佳方案（前100）"
        elif which == "abs":
            data = getattr(self, "_last_abs_list", [])
            title = "最高战力方案（前100）"
        elif which == "bs":
            data = getattr(self, "_last_bs_list", [])
            title = "职业推荐方案（前100）"
        else:
            from tkinter import messagebox
            messagebox.showerror("错误", f"未知导出类型：{which}")
            return

        if not data:
            from tkinter import messagebox
            messagebox.showwarning("提示", "当前没有可导出的结果。请先计算。")
            return

        # 3) 生成文档
        doc = Document()
        doc.add_heading(title, level=1)

        def lv6_attrs(levels: dict) -> str:
            arr = [a for a, lv in levels.items() if lv == 6]
            return "，".join(arr) if arr else "（无）"

        # 统一解包：兼容 ModuleSolution 与老 pack
        def _unpack(item):
            # 新：ModuleSolution
            if hasattr(item, "modules") and hasattr(item, "attr_breakdown"):
                mods = list(item.modules)
                pts  = dict(item.attr_breakdown)
                levels = {a: points_to_level(v) for a, v in pts.items()}
                c6    = sum(1 for v in levels.values() if v == 6)
                sumlv = sum(levels.values())
                waste = sum(max(0, v - 20) for v in pts.values())
                total_points = sum(pts.values())
                return mods, pts, levels, c6, sumlv, waste, total_points
            # 旧：pack 元组
            _, _, combo_gems, pts, levels, c6, sumlv, waste, total_points = item
            return list(combo_gems), dict(pts), dict(levels), c6, sumlv, waste, total_points

        for k, item in enumerate(data, start=1):
            combo_gems, pts, levels, c6, sumlv, waste, total_points = _unpack(item)

            doc.add_heading(f"方案 #{k}", level=2)

            # 使用模组：逐行显示
            for g in combo_gems:
                doc.add_paragraph(str(format_gem_str(g)))

            # 概览与统计
            doc.add_paragraph("等级概览：" + levels_pretty(levels))
            doc.add_paragraph("6级属性：" + lv6_attrs(levels))
            doc.add_paragraph(f"统计：6级数量={c6}，等级总和={sumlv}，浪费点={waste}，点数总和={total_points}")

        # 4) 保存对话框
        from tkinter import filedialog, messagebox
        path = filedialog.asksaveasfilename(
            title="导出为 Word",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile=title.replace("（前100）", "").strip() + ".docx"
        )
        if not path:
            return

        # 5) 保存
        try:
            doc.save(path)
        except Exception as e:
            messagebox.showerror("保存失败", f"无法写入文件：\n{e}")
            return

        messagebox.showinfo("完成", f"已导出到：\n{path}")

    def delete_low_score_gems(self):
        import tkinter.simpledialog as sd
        # 输入阈值（包含第三词条，最高 30）
        x = sd.askinteger("删除低级模组", "请输入点数阈值（最大30）：", minvalue=0, maxvalue=30)
        if x is None:
            return

        # 统计拟删除：p1+p2+p3（p3 可能不存在，按 0 处理）
        cand = [
            i for i, g in enumerate(self.gems)
            if (g.get("p1", 0) + g.get("p2", 0) + g.get("p3", 0)) < x
        ]

        if not cand:
            messagebox.showinfo("提示", f"没有点数和 < {x} 的模组。")
            return

        if not messagebox.askyesno("确认", f"将删除 {len(cand)} 个模组（点数和 < {x}，已包含第三词条）。是否继续？"):
            return

        # 从后往前删，避免索引位移
        for i in sorted(cand, reverse=True):
            del self.gems[i]

        # 重建列表（一次性清空更快）
        self.tree.delete(*self.tree.get_children())
        for g in self.gems:
            self.tree.insert("", tk.END, values=(format_gem_str(g),))

        messagebox.showinfo("完成", f"已删除 {len(cand)} 个模组。")


    # ——互斥联动：当任一目标/排除改动时调用
    def _entry_changed(self, _entry):
        self._refresh_mutex_states(show_prompt=True)

    # 计算当前已选目标/排除集合，刷新禁用项，并对冲突给出提示/清空
    def _refresh_mutex_states(self, show_prompt=False):
        targets_set = set()
        for ent, sp in self.targets:
            v = ent.get_value()
            if v and v in ATTRS:
                targets_set.add(v)
        excludes_set = set()
        for ent in self.exclude_entries:
            v = ent.get_value()
            if v and v in ATTRS:
                excludes_set.add(v)

        # 互斥：目标禁用“排除已选”，排除禁用“目标已选”
        for ent, _ in self.targets:
            ent.set_disabled(excludes_set)
        for ent in self.exclude_entries:
            ent.set_disabled(targets_set)

        # 冲突时清空排除侧的冲突项（也可以改为清空目标侧）
        overlap = targets_set & excludes_set
        if overlap:
            if show_prompt:
                messagebox.showwarning("冲突", f"目标属性与排除属性不能相同：{', '.join(sorted(overlap))}")
            for ent in self.exclude_entries:
                if ent.get_value() in overlap:
                    ent.var.set("")
            self.after(10, lambda: self._refresh_mutex_states(False))

    def build_ui(self):
        # 外层左右两栏
        root = ttk.Frame(self, padding=12)
        root.pack(fill=tk.BOTH, expand=True)

        left = ttk.Frame(root)
        right = ttk.Frame(root)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0,8))
        right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # 左侧：添加模组
        lf = ttk.LabelFrame(left, text="添加模组")
        lf.pack(fill=tk.X, pady=(0,8))
        row = 0
        ttk.Label(lf, text="属性1").grid(row=row, column=0, sticky="w", padx=6, pady=6)
        self.a1 = AttrPickerField(lf, ATTRS, width=12, open_handler=lambda e: self._open_group_picker_addgem())
        self.a1.grid(row=row, column=1, sticky="w", padx=6, pady=6)

        ttk.Label(lf, text="点数").grid(row=row, column=2, sticky="w", padx=6)
        self.p1 = tk.Spinbox(lf, from_=POINT_MIN, to=POINT_MAX, width=5)
        self.p1.grid(row=row, column=3, sticky="w", padx=6)

        row += 1
        ttk.Label(lf, text="属性2").grid(row=row, column=0, sticky="w", padx=6, pady=6)
        self.a2 = AttrPickerField(lf, ATTRS, width=12, open_handler=lambda e: self._open_group_picker_addgem())
        self.a2.grid(row=row, column=1, sticky="w", padx=6, pady=6)

        ttk.Label(lf, text="点数").grid(row=row, column=2, sticky="w", padx=6)
        self.p2 = tk.Spinbox(lf, from_=POINT_MIN, to=POINT_MAX, width=5)
        self.p2.grid(row=row, column=3, sticky="w", padx=6)

        row += 1
        ttk.Label(lf, text="属性3").grid(row=row, column=0, sticky="w", padx=6, pady=6)
        self.a3 = AttrPickerField(lf, ATTRS, width=12, open_handler=lambda e: self._open_group_picker_addgem())
        self.a3.grid(row=row, column=1, sticky="w", padx=6, pady=6)

        ttk.Label(lf, text="点数").grid(row=row, column=2, sticky="w", padx=6)
        self.p3 = tk.Spinbox(lf, from_=POINT_MIN, to=POINT_MAX, width=5)
        self.p3.grid(row=row, column=3, sticky="w", padx=6)

        ttk.Button(lf, text="添加到列表", command=self.add_gem).grid(row=0, column=4, rowspan=3, padx=10)

        # 左侧：工具条 导出导入
        tool = ttk.Frame(left)
        tool.pack(fill=tk.X, pady=(0,6))
        ttk.Button(tool, text="导入Excel", command=self.import_excel).pack(side=tk.LEFT)
        ttk.Button(tool, text="导出Excel", command=self.export_excel).pack(side=tk.LEFT, padx=6)
        ttk.Button(tool, text="导出模板", command=self.export_template).pack(side=tk.LEFT, padx=6)
        ttk.Button(tool, text="识别截图", command=self.import_from_image).pack(side=tk.LEFT, padx=6)
        ttk.Button(tool, text="识别文件夹", command=self.import_from_folder).pack(side=tk.LEFT, padx=6)
        # 左侧：工具条 抓包
        tool2 = ttk.Frame(left)
        tool2.pack(fill=tk.X, pady=(6,0))
        self.capture_btn = ttk.Button(tool2, text="开始抓包", command=self.toggle_packet_capture)
        self.capture_btn.pack(side=tk.LEFT)
        self.capture_status = ttk.Label(tool2, text="未开始")
        self.capture_status.pack(side=tk.LEFT, padx=(10,0))

        # 左侧：职业 / 期望 / 排除 / 计算
        targetf = ttk.LabelFrame(left, text="目标词条和等级（可不填）")
        targetf.pack(fill=tk.X, pady=(6,6))

        # 职业选择（点击文本框即弹窗）
        jobf = ttk.Frame(left)
        jobf.pack(fill=tk.X, pady=(0,6))
        ttk.Label(jobf, text="职业：").pack(side=tk.LEFT, padx=(6,2))

        self.job_field = AttrPickerField(
            jobf, CLASS_LIST, width=12,
            open_handler=lambda e: self._open_group_picker_job()
        )
        self.job_field.pack(side=tk.LEFT, padx=4)

        targets_container = ttk.Frame(targetf)
        targets_container.pack(fill=tk.X, padx=6, pady=4)
        
        left_targets = ttk.Frame(targets_container)
        left_targets.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        right_targets = ttk.Frame(targets_container)
        right_targets.pack(side=tk.LEFT, fill=tk.X, expand=True)
    
        for i in range(3):
            rowf = ttk.Frame(left_targets)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"目标{i+1}").pack(side=tk.LEFT, anchor="w")
            ent = AttrPickerField(rowf, ATTRS, on_change=self._entry_changed, width=10, open_handler=lambda e: self._open_group_picker_targets())
            ent.pack(side=tk.LEFT, padx=(6,8))
            ttk.Label(rowf, text="目标等级").pack(side=tk.LEFT)
            sp = tk.Spinbox(rowf, from_=0, to=6, width=5)
            sp.delete(0, tk.END); sp.insert(0, "0")
            sp.pack(side=tk.LEFT, padx=6)
            self.targets.append((ent, sp))
            self.target_entries.append(ent)

        for i in range(3, 6):
            rowf = ttk.Frame(right_targets)
            rowf.pack(fill=tk.X, pady=4)
            ttk.Label(rowf, text=f"目标{i+1}").pack(side=tk.LEFT, anchor="w")
            ent = AttrPickerField(rowf, ATTRS, on_change=self._entry_changed, width=10, open_handler=lambda e: self._open_group_picker_targets())
            ent.pack(side=tk.LEFT, padx=(6,8))
            ttk.Label(rowf, text="目标等级").pack(side=tk.LEFT)
            sp = tk.Spinbox(rowf, from_=0, to=6, width=5)
            sp.delete(0, tk.END); sp.insert(0, "0")
            sp.pack(side=tk.LEFT, padx=6)
            self.targets.append((ent, sp))
            self.target_entries.append(ent)
        
        exf = ttk.LabelFrame(left, text="排除词条（可不填）")
        exf.pack(fill=tk.X, pady=(0,6))
        labels = [f"排除{i}" for i in range(1, 10)]
        idx = 0
        for row_i in range(3):
            rowf = ttk.Frame(exf)
            rowf.pack(fill=tk.X, padx=6, pady=3)
            for _ in range(3):
                ttk.Label(rowf, text=labels[idx]).pack(side=tk.LEFT, padx=(0,4))
                ent = AttrPickerField(rowf, ATTRS, on_change=self._entry_changed, width=12, open_handler=lambda e: self._open_group_picker_excludes())
                ent.pack(side=tk.LEFT, padx=(0,10))
                self.exclude_entries.append(ent)
                idx += 1
        modef = ttk.LabelFrame(left, text="计算方式", padding=(4, 2)) 
        modef.pack(fill=tk.X, pady=(2, 2)) 

        style = ttk.Style()
        style.configure("Compact.TRadiobutton", padding=(2, 0))
        style.configure("Compact.TButton", padding=(6, 2))

        self.compute_mode = tk.StringVar(value="fast")
        rb_fast = ttk.Radiobutton(modef, text="快速", value="fast", variable=self.compute_mode, style="Compact.TRadiobutton")
        rb_full = ttk.Radiobutton(modef, text="全部", value="full", variable=self.compute_mode, style="Compact.TRadiobutton")
        btn_calc = ttk.Button(modef, text="计算组合", command=self.compute, style="Compact.TButton")

        # 横排三列
        modef.columnconfigure(0, weight=1)
        modef.columnconfigure(1, weight=1)
        modef.columnconfigure(2, weight=0)

        rb_fast.grid(row=0, column=0, sticky="w", padx=(6, 2), pady=2)
        rb_full.grid(row=0, column=1, sticky="w", padx=(2, 4), pady=2)
        btn_calc.grid(row=0, column=2, sticky="e", padx=(4, 6), pady=2)

        # 左侧：已添加的模组
        listf = ttk.LabelFrame(left, text="已添加的模组")
        listf.pack(fill=tk.BOTH, expand=True)
        self.tree = ttk.Treeview(listf, columns=("desc",), show="headings", height=9)
        self.tree.heading("desc", text="模组属性")
        self.tree.column("desc", anchor=tk.CENTER, width=460, stretch=True)
        self.tree.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        btns = ttk.Frame(listf)
        btns.pack(fill=tk.X, padx=6, pady=(0,8))
        ttk.Button(btns, text="删除选中", command=self.delete_selected).pack(side=tk.LEFT)
        ttk.Button(btns, text="清空列表", command=self.clear_all).pack(side=tk.LEFT, padx=8)
        ttk.Button(btns, text="删除低级模组", command=self.delete_low_score_gems).pack(side=tk.LEFT, padx=8)

            # 右侧：计算结果（统一大小 + 滚动条）
        resf = ttk.LabelFrame(right, text="计算结果")
        resf.pack(fill=tk.BOTH, expand=True)

        # 导出按钮放在最上方，始终可见
        exportf = ttk.Frame(resf)
        exportf.pack(fill=tk.X, padx=8, pady=(8,8))
        ttk.Button(exportf, text="导出需求最佳（Word）",
                command=lambda: self.export_solutions_to_docx("req")).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(exportf, text="导出最高等级（Word）",
                command=lambda: self.export_solutions_to_docx("abs")).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(exportf, text="导出战力推荐（Word）",
                command=lambda: self.export_solutions_to_docx("bs")).pack(side=tk.LEFT, padx=(0,8))

        def _make_scrolled_text(parent, title, height=12):
            # 标题
            ttk.Label(parent, text=title, font=("Segoe UI", 10, "bold")).pack(
                anchor="w", padx=8, pady=(4,2)
            )
            # 带滚动条的 Text
            box = ttk.Frame(parent)
            box.pack(fill=tk.BOTH, expand=True, padx=8, pady=(0,6))
            txt = tk.Text(box, height=12, wrap="word", state="disabled")  # 添加 state="disabled"
            vsb = ttk.Scrollbar(box, orient="vertical", command=txt.yview)
            txt.configure(yscrollcommand=vsb.set)
            txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            vsb.pack(side=tk.RIGHT, fill=tk.Y)
            return txt

        # 三个方案：大小一致 + 各自滚动条
        self.abs_text = _make_scrolled_text(resf, "战力最高方案", height=12)
        self.req_text = _make_scrolled_text(resf, "需求最佳方案(仅在有目标或排除属性时生效)", height=12)
        self.bs_text  = _make_scrolled_text(resf, "职业推荐方案(仅在选择职业时有效)", height=12)

        # 初始化互斥状态
        self._refresh_mutex_states(False)

    # ---------- 列表操作 ----------
    def add_gem(self):
        a1 = self.a1.get_value()
        a2 = self.a2.get_value()
        a3 = self.a3.get_value()
        try:
            p1 = int(self.p1.get())
            p2 = int(self.p2.get())
            p3 = int(self.p3.get())
        except ValueError:
            messagebox.showerror("错误", "点数必须是整数。")
            return
        names = [x for x in [a1, a2, a3] if x]
        if len(names) < 2:
            messagebox.showerror("错误", "至少需要选择两个词条。")
            return
        if a1 not in ATTRS or a2 not in ATTRS:
            messagebox.showerror("错误", "属性必须从候选中选择或拼写正确。")
            return
        if a1 == a2:
            messagebox.showerror("错误", "同一颗模组的两条属性不能相同。")
            return
        if not (POINT_MIN <= p1 <= POINT_MAX and POINT_MIN <= p2 <= POINT_MAX):
            messagebox.showerror("错误", f"点数范围应为 {POINT_MIN}~{POINT_MAX}。")
            return
        g = {"a1": a1, "p1": p1, "a2": a2, "p2": p2}
        if a3:
            g["a3"] = a3
            g["p3"] = p3
        self.gems.append(g)
        self.tree.insert("", tk.END, values=(format_gem_str(g),))
        self.a1.var.set(""); self.a2.var.set(""); self.a3.var.set("")
        self.p1.delete(0, tk.END); self.p1.insert(0, str(POINT_MIN))
        self.p2.delete(0, tk.END); self.p2.insert(0, str(POINT_MIN))
        self.p3.delete(0, tk.END); self.p3.insert(0, str(POINT_MIN))

    def delete_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        idxs = sorted([self.tree.index(i) for i in sel], reverse=True)
        for i in idxs:
            del self.gems[i]
        for i in sel:
            self.tree.delete(i)

    def clear_all(self):
        if messagebox.askyesno("确认", "确定清空所有模组？"):
            self.gems.clear()
            for i in self.tree.get_children():
                self.tree.delete(i)

    # ---------- Excel 导入导出 ----------
    def _require_pandas(self):
        if not HAS_PANDAS:
            messagebox.showerror("缺少依赖", "需要安装 pandas 和 openpyxl 才能读写 Excel：\n\npip install pandas openpyxl")
            return False
        return True

    def export_template(self):
        if not self._require_pandas():
            return
        path = filedialog.asksaveasfilename(
            title="导出模板为...",
            defaultextension=".xlsx",
            filetypes=[("Excel 文件", "*.xlsx")]
        )
        if not path:
            return
        # 支持三个词条的模板
        df = pd.DataFrame([
            {"属性1": "", "点数1": "", "属性2": "", "点数2": "", "属性3": "", "点数3": ""}
        ])
        df.to_excel(Path(path), index=False)
        messagebox.showinfo("完成", "模板已导出。填写后可用\"导入Excel\"批量载入。\n注意：属性3和点数3为可选项，可以留空。")

    def export_excel(self):
        if not self._require_pandas():
            return
        if not self.gems:
            messagebox.showwarning("提示", "当前没有可导出的模组。")
            return
        
        path = filedialog.asksaveasfilename(
            title="导出为...",
            defaultextension=".xlsx",
            filetypes=[("Excel 文件", "*.xlsx")]
        )
        if not path:
            return
        
        # Ensure the file has .xlsx extension
        if not path.lower().endswith('.xlsx'):
            path += '.xlsx'
        
        # 支持三个词条的导出
        rows = []
        for g in self.gems:
            row = {
                "属性1": g["a1"], 
                "点数1": g["p1"], 
                "属性2": g["a2"], 
                "点数2": g["p2"],
                "属性3": g.get("a3", ""),  # 可能不存在第三个词条
                "点数3": g.get("p3", "")
            }
            rows.append(row)
        
        df = pd.DataFrame(rows, columns=["属性1","点数1","属性2","点数2","属性3","点数3"])
        
        try:
            # Use string path directly instead of Path object
            df.to_excel(path, index=False)
            messagebox.showinfo("完成", f"已导出到：\n{path}")
        except PermissionError:
            messagebox.showerror("导出失败", "文件可能正在被其他程序使用，请关闭后重试。")
        except Exception as e:
            messagebox.showerror("导出失败", f"无法写入文件：\n{str(e)}")
        
    def import_excel(self):
        if not self._require_pandas():
            return
        path = filedialog.askopenfilename(
            title="选择 Excel 文件",
            filetypes=[("Excel 文件", "*.xlsx *.xls")]
        )
        if not path:
            return
        try:
            df = pd.read_excel(Path(path), dtype={
                "属性1": str, "点数1": "Int64", 
                "属性2": str, "点数2": "Int64",
                "属性3": str, "点数3": "Int64"
            })
        except Exception as e:
            messagebox.showerror("读取失败", f"无法读取文件：\n{e}")
            return
        
        # 检查必需列
        required_cols = {"属性1","点数1","属性2","点数2"}
        if not required_cols.issubset(set(df.columns)):
            messagebox.showerror("格式错误", "缺少列：属性1、点数1、属性2、点数2\n注意：属性3、点数3为可选列")
            return
        
        loaded = 0
        errors = []
        
        for idx, row in df.iterrows():
            # 处理必需的前两个词条
            a1 = str(row["属性1"]).strip() if pd.notna(row["属性1"]) else ""
            a2 = str(row["属性2"]).strip() if pd.notna(row["属性2"]) else ""
            
            # 处理可选的第三个词条
            a3 = ""
            p3 = None
            if "属性3" in df.columns and "点数3" in df.columns:
                a3 = str(row["属性3"]).strip() if pd.notna(row["属性3"]) else ""
                try:
                    p3 = int(row["点数3"]) if pd.notna(row["点数3"]) and str(row["点数3"]).strip() else None
                except:
                    p3 = None
            
            try:
                p1 = int(row["点数1"]) if pd.notna(row["点数1"]) else None
                p2 = int(row["点数2"]) if pd.notna(row["点数2"]) else None
            except Exception:
                p1 = p2 = None
            
            row_no = idx + 2
            
            # 校验必需词条
            if a1 not in ATTRS or a2 not in ATTRS:
                errors.append(f"第{row_no}行：前两个属性名不在列表中")
                continue
            if a1 == a2:
                errors.append(f"第{row_no}行：同一模组前两条属性重复")
                continue
            if p1 is None or p2 is None:
                errors.append(f"第{row_no}行：前两个点数缺失或不是整数")
                continue
            if not (POINT_MIN <= p1 <= POINT_MAX and POINT_MIN <= p2 <= POINT_MAX):
                errors.append(f"第{row_no}行：前两个点数需在 {POINT_MIN}~{POINT_MAX}")
                continue
            
            # 校验可选的第三个词条
            if a3:  # 如果有第三个属性
                if a3 not in ATTRS:
                    errors.append(f"第{row_no}行：第三个属性名不在列表中")
                    continue
                if a3 in [a1, a2]:
                    errors.append(f"第{row_no}行：第三个属性与前面的属性重复")
                    continue
                if p3 is None:
                    errors.append(f"第{row_no}行：第三个属性有值但点数缺失")
                    continue
                if not (POINT_MIN <= p3 <= POINT_MAX):
                    errors.append(f"第{row_no}行：第三个点数需在 {POINT_MIN}~{POINT_MAX}")
                    continue
            
            # 创建模组字典
            g = {"a1": a1, "p1": p1, "a2": a2, "p2": p2}
            if a3 and p3 is not None:
                g["a3"] = a3
                g["p3"] = p3
            
            self.gems.append(g)
            self.tree.insert("", tk.END, values=(format_gem_str(g),))
            loaded += 1
        
        if errors:
            msg = "部分行导入失败：\n" + "\n".join(errors[:20])
            if len(errors) > 20:
                msg += f"\n... 还有 {len(errors)-20} 条"
            messagebox.showwarning("完成（有错误）", f"成功导入 {loaded} 条。\n\n{msg}")
        else:
            messagebox.showinfo("完成", f"成功导入 {loaded} 条。")

    # ---------- 图像识别（单图/文件夹） ----------
    def import_from_image(self):
        path = filedialog.askopenfilename(
            title="选择一张包含模组列表的截图",
            filetypes=[("图片","*.png *.jpg *.jpeg *.bmp *.webp")]
        )
        if not path:
            return
        detect_t, classify_t = _load_icon_templates()
        if not detect_t:
            messagebox.showerror("缺少图标", f"未在 {ICON_DIR} 找到模板图标，请先放入 icons/属性名.png")
            return
        try:
            gems = parse_screenshot_to_gems(Path(path), detect_t, classify_t)
        except Exception as e:
            messagebox.showerror("识别失败", str(e))
            return
        added = 0
        for g in gems:
            if g["a1"] in ATTRS and g["a2"] in ATTRS:
                self.gems.append(g)
                self.tree.insert("", tk.END, values=(format_gem_str(g),))
                added += 1
        messagebox.showinfo("完成", f"已识别添加 {added} 条。")

    def import_from_folder(self):
        # 1) 选择目录（主线程）
        folder = filedialog.askdirectory(title="选择包含截图的文件夹")
        if not folder:
            return
        folder = Path(folder)

        # 2) 加载模板（主线程）
        detect_t, classify_t = _load_icon_templates()
        if not detect_t:
            messagebox.showerror("缺少图标", f"未在 {ICON_DIR} 找到模板图标，请先放入 icons/属性名.png")
            return

        # 3) 创建进度弹窗（主线程）
        win = tk.Toplevel(self)
        win.title("OCR 进度")
        win.resizable(False, False)
        ttk.Label(win, text="正在识别截图…").pack(padx=12, pady=(12, 6))
        pb = ttk.Progressbar(win, length=360, mode="determinate", maximum=1)
        pb.pack(padx=12, pady=(0, 4))
        lab = ttk.Label(win, text="")
        lab.pack(padx=12, pady=(0, 12))
        win.update_idletasks()
        try:
            x = self.winfo_rootx() + (self.winfo_width()-win.winfo_width())//2
            y = self.winfo_rooty() + (self.winfo_height()-win.winfo_height())//2
            win.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

        # 4) 线程通信队列（主线程）
        q = queue.Queue()

        # 5) 主线程：轮询队列刷新 UI
        def poll():
            try:
                while True:
                    item = q.get_nowait()
                    kind = item[0]
                    if kind == "progress":
                        _, i, total, name = item
                        pb["maximum"] = max(1, total)
                        pb["value"] = i
                        lab.config(text=f"{i}/{total}  {name}")
                    elif kind == "result":
                        _, gems, file_cnt, fails, err = item
                        try:
                            win.destroy()
                        except Exception:
                            pass
                        if err:
                            messagebox.showerror("识别失败", err)
                            return
                        # 写回结果到主界面
                        added = 0
                        for g in gems:
                            if g.get("a1") in ATTRS and g.get("a2") in ATTRS:
                                self.gems.append(g)
                                self.tree.insert("", tk.END, values=(format_gem_str(g),))
                                added += 1
                        msg = [f"共扫描图片：{file_cnt} 张", f"成功解析模组：{added} 条"]
                        if fails:
                            msg.append(f"失败图片：{len(fails)} 张（仅展示前5条）")
                            for p, e in fails[:5]:
                                name = getattr(p, "name", str(p))
                                msg.append(f"- {name}：{e}")
                            if len(fails) > 5:
                                msg.append("...")
                        messagebox.showinfo("识别完成", "\n".join(msg))
                        return
            except queue.Empty:
                pass
            self.after(50, poll)  # 继续轮询

        def worker():
            try:
                def _cb(i: int, total: int, name: str):
                    q.put(("progress", i, total, name))
                try:
                    gems, file_cnt, fails = parse_folder_to_gems(folder, detect_t, classify_t, progress_cb=_cb)
                except TypeError:
                    gems, file_cnt, fails = parse_folder_to_gems(folder, detect_t, classify_t)
                q.put(("result", gems, file_cnt, fails, None))
            except Exception as e:
                q.put(("result", [], 0, [], str(e)))

        threading.Thread(target=worker, daemon=True).start()
        poll()


    # 计算
    def compute(self):
        slots = 4
        gems_all = list(self.gems)

        # 读取排除属性
        exclude_set = {ent.get_value() for ent in self.exclude_entries if ent.get_value()}

        if len(gems_all) < slots:
            messagebox.showerror("错误", f"需要至少 {slots} 颗模组。")
            return

        # 导入外部模块
        if ModuleOptimizerCalculations is None:
            messagebox.showerror("缺少计算模块", "未能导入外部计算模块（module_compute.py / module_optimizer_calculations.py）。")
            return

        # 先过滤掉带有排除属性的模组
        def _gem_has_excluded(g):
            return any(x in exclude_set for x in [g.get("a1"), g.get("a2"), g.get("a3")])
        
        filtered_gems = [g for g in gems_all if not _gem_has_excluded(g)]
        
        if len(filtered_gems) < slots:
            messagebox.showwarning("无可用组合", "排除属性过多或模组过少，无法构成组合。")
            return

        # 解析目标需求
        target_requirements = self._parse_target_requirements()

        # 获取计算模式
        mode = getattr(self, "compute_mode", None).get() if hasattr(self, "compute_mode") else "fast"

        # 清空三个面板
        for box in (self.req_text, self.abs_text, self.bs_text):
            box.configure(state="normal")
            box.delete("1.0", tk.END)
            box.configure(state="disabled")

        if mode == "fast":
            # 快速模式：同步执行
            self._compute_fast_sync_with_requirements(filtered_gems, exclude_set, slots, target_requirements)
        else:
            # 全部模式：异步执行
            self._compute_full_async_with_requirements(filtered_gems, exclude_set, slots, target_requirements)

    def _compute_full_async_with_requirements(self, filtered_gems, exclude_set, slots, target_requirements):
        """异步全部计算：支持需求最佳方案的版本"""
        import threading
        import queue
        
        self._cancel_flag = threading.Event()
        self._compute_queue = queue.Queue()
        
        # 创建进度窗口
        self._progress_window = tk.Toplevel(self)
        self._progress_window.title("正在计算...")
        self._progress_window.resizable(False, False)
        self._progress_window.grab_set()
        self._progress_window.protocol("WM_DELETE_WINDOW", self._cancel_computation)
        
        # 进度UI
        frame = ttk.Frame(self._progress_window, padding=10)
        frame.pack(fill=tk.BOTH, expand=True)
        
        self._progress_var = tk.StringVar(value="准备中...")
        ttk.Label(frame, textvariable=self._progress_var).pack(pady=(0,10))
        
        self._progress_bar = ttk.Progressbar(frame, length=400, mode="determinate", maximum=100)
        self._progress_bar.pack(pady=(0,10))
        
        ttk.Button(frame, text="取消", command=self._cancel_computation).pack()
        
        # 居中显示
        self._progress_window.update_idletasks()
        try:
            x = self.winfo_rootx() + (self.winfo_width()-self._progress_window.winfo_width())//2
            y = self.winfo_rooty() + (self.winfo_height()-self._progress_window.winfo_height())//2
            self._progress_window.geometry(f"+{max(0,x)}+{max(0,y)}")
        except:
            pass
        
        # 启动工作线程 - 使用修正后的双Top-K版本
        worker = threading.Thread(
            target=self._compute_full_worker_with_requirements_dual_topk, 
            args=(filtered_gems, exclude_set, slots, target_requirements),
            daemon=True
        )
        worker.start()
        
        # 开始轮询进度
        self._poll_compute_progress_with_requirements()

    def _compute_full_worker_with_requirements_dual_topk(self, filtered_gems, exclude_set, slots, target_requirements):
        """
        全部计算工作线程：支持双Top-K的版本，分别生成战力和需求两套结果
        """
        try:
            # 导入配置（同原版）
            try:
                from module_types import (
                    ATTR_THRESHOLDS, BASIC_ATTR_POWER_MAP, SPECIAL_ATTR_POWER_MAP,
                    TOTAL_ATTR_POWER_MAP, ATTR_NAME_TYPE_MAP
                )
            except ImportError:
                ATTR_THRESHOLDS = [1, 4, 8, 12, 16, 20]
                BASIC_ATTR_POWER_MAP = {1: 7, 2: 14, 3: 29, 4: 44, 5: 167, 6: 254}
                SPECIAL_ATTR_POWER_MAP = {1: 14, 2: 29, 3: 59, 4: 89, 5: 298, 6: 448}
                TOTAL_ATTR_POWER_MAP = {i: int(i * 0.1) for i in range(0, 201)}
                ATTR_NAME_TYPE_MAP = {}
            
            opt = ModuleOptimizerCalculations(
                exclude_set=exclude_set,
                slots=slots,
                thresholds=ATTR_THRESHOLDS,
                attr_type_map=ATTR_NAME_TYPE_MAP,
                basic_attr_power_map=BASIC_ATTR_POWER_MAP,
                special_attr_power_map=SPECIAL_ATTR_POWER_MAP,
                total_attr_power_map=TOTAL_ATTR_POWER_MAP
            )
            
            # 进度回调
            def enum_progress_callback(current, total, message=""):
                if not self._cancel_flag.is_set():
                    progress = int(80 * current / total) if total > 0 else 0
                    self._compute_queue.put(("progress", progress, 100, message))
            
            from math import comb   
            total_combinations = comb(len(filtered_gems), slots)
            
            # 使用双Top-K枚举方法
            if total_combinations > 50000:  # 大数据集使用多线程版本
                battle_solutions, requirement_solutions = opt.enum_all_combinations_ui_threaded_memory_safe_dual_topk(
                    filtered_gems,
                    slots=slots,
                    progress_callback=enum_progress_callback,
                    cancel_flag=self._cancel_flag,
                    max_keep=500,
                    target_requirements=target_requirements
                )
            else:
                battle_solutions, requirement_solutions = opt.enum_all_combinations_ui_memory_safe_dual_topk(
                    filtered_gems,
                    slots=slots,
                    progress_callback=enum_progress_callback,
                    max_keep=500,
                    target_requirements=target_requirements
                )

            if self._cancel_flag.is_set():
                return

            self._compute_queue.put(("progress", 80, 100, "生成最终方案..."))
            
            # 直接使用双Top-K的结果，取前30个
            abs_list = battle_solutions[:30]
            req_list = requirement_solutions[:30]
            
            # 职业推荐（目前使用战力排序）
            bs_list = abs_list
            
            self._compute_queue.put(("progress", 100, 100, "完成"))
            self._compute_queue.put(("result", req_list, abs_list, bs_list, opt, target_requirements))

        except Exception as e:
            import traceback
            traceback.print_exc()
            self._compute_queue.put(("error", str(e)))

    def _parse_target_requirements(self) -> List[Tuple[str, int]]:
        """解析用户设置的目标属性和等级，只返回有效的目标需求"""
        requirements = []
        for ent, sp in self.targets:
            attr_name = ent.get_value()
            if attr_name and attr_name.strip():
                try:
                    target_level = int(sp.get())
                    if target_level > 0:  # 只考虑大于0的等级要求
                        requirements.append((attr_name, target_level))
                    # 移除下面这行：如果等级为0就不算有目标需求
                    # else:
                    #     requirements.append((attr_name, 1))  # 默认至少要1级
                except ValueError:
                    # 解析失败，跳过这个目标
                    continue
        return requirements

    def _compute_fast_sync_with_requirements(self, filtered_gems, exclude_set, slots, target_requirements):
        """快速计算：支持需求最佳方案的专门评分"""
        try:
            # 简单进度条
            dlg = _Progress(self, 100, on_cancel=lambda: self._set_cancel(dlg))
            self._cancel = False
            
            # 导入正确的配置数据
            try:
                from module_types import (
                    ATTR_THRESHOLDS, BASIC_ATTR_POWER_MAP, SPECIAL_ATTR_POWER_MAP,
                    TOTAL_ATTR_POWER_MAP, ATTR_NAME_TYPE_MAP
                )
            except ImportError:
                ATTR_THRESHOLDS = [1, 4, 8, 12, 16, 20]
                BASIC_ATTR_POWER_MAP = {1: 7, 2: 14, 3: 29, 4: 44, 5: 167, 6: 254}
                SPECIAL_ATTR_POWER_MAP = {1: 14, 2: 29, 3: 59, 4: 89, 5: 298, 6: 448}
                TOTAL_ATTR_POWER_MAP = {i: int(i * 0.1) for i in range(0, 201)}
                ATTR_NAME_TYPE_MAP = {}
            
            opt = ModuleOptimizerCalculations(
                exclude_set=exclude_set, 
                slots=slots,
                thresholds=ATTR_THRESHOLDS,
                attr_type_map=ATTR_NAME_TYPE_MAP,
                basic_attr_power_map=BASIC_ATTR_POWER_MAP,
                special_attr_power_map=SPECIAL_ATTR_POWER_MAP,
                total_attr_power_map=TOTAL_ATTR_POWER_MAP
            )
            
            # 预过滤候选模组
            candidates = opt.prefilter_ui_mods(filtered_gems)
            dlg.set_progress(20, 100)
            
            if not candidates or len(candidates) < slots:
                dlg.destroy()
                self._render_empty_results("（快速）")
                return

            if self._cancel:
                dlg.destroy()
                return
            
            # 使用快速多解方法生成更多候选方案
            dlg.set_progress(30, 100)
            solutions = opt.fast_multi_solutions_ui(
                candidates, 
                slots=slots, 
                prioritize_high_level=True, 
                num_solutions=200,  # 生成更多解以便后续排序
                max_attempts=500
            )
            
            if self._cancel:
                dlg.destroy()
                return
                
            dlg.set_progress(70, 100)
            
            if solutions:
            # 分别计算两种排序
                
                # 1. 按战力分排序（用于战力最高方案）
                battle_power_sorted = []
                for sol in solutions:
                    bp = opt.combat_power_from_attr_breakdown(sol.attr_breakdown)
                    battle_power_sorted.append((bp, sol))
                battle_power_sorted.sort(key=lambda x: x[0], reverse=True)
                abs_list = [sol for bp, sol in battle_power_sorted[:30]]
                
                # 2. 按需求评分排序（仅当有具体目标需求时）
                if target_requirements and len(target_requirements) > 0:
                    requirement_scored = []
                    for sol in solutions:
                        req_score = opt.requirement_score_from_attr_breakdown(sol.attr_breakdown, target_requirements)
                        requirement_scored.append((req_score, sol))
                    requirement_scored.sort(key=lambda x: x[0], reverse=True)
                    req_list = [sol for score, sol in requirement_scored[:30]]
                else:
                    req_list = abs_list
                
                # 3. 职业推荐方案（暂时也使用战力排序，后续根据职业属性优化）
                bs_list = abs_list
                
            else:
                req_list = abs_list = bs_list = []
            
            dlg.set_progress(100, 100)
            dlg.destroy()
            
            # 渲染结果
            self._render_results_with_requirements(req_list, abs_list, bs_list, "（快速）", opt, target_requirements)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("计算错误", f"快速计算发生错误：{e}")
            try:
                dlg.destroy()
            except:
                pass

    
    def _compute_full_worker_with_requirements(self, filtered_gems, exclude_set, slots, target_requirements):
        """全部计算工作线程：支持需求最佳方案评分"""
        try:
            # 导入配置（同快速计算）
            try:
                from module_types import (
                    ATTR_THRESHOLDS, BASIC_ATTR_POWER_MAP, SPECIAL_ATTR_POWER_MAP,
                    TOTAL_ATTR_POWER_MAP, ATTR_NAME_TYPE_MAP
                )
            except ImportError:
                ATTR_THRESHOLDS = [1, 4, 8, 12, 16, 20]
                BASIC_ATTR_POWER_MAP = {1: 7, 2: 14, 3: 29, 4: 44, 5: 167, 6: 254}
                SPECIAL_ATTR_POWER_MAP = {1: 14, 2: 29, 3: 59, 4: 89, 5: 298, 6: 448}
                TOTAL_ATTR_POWER_MAP = {i: int(i * 0.1) for i in range(0, 201)}
                ATTR_NAME_TYPE_MAP = {}
            
            opt = ModuleOptimizerCalculations(
                exclude_set=exclude_set,
                slots=slots,
                thresholds=ATTR_THRESHOLDS,
                attr_type_map=ATTR_NAME_TYPE_MAP,
                basic_attr_power_map=BASIC_ATTR_POWER_MAP,
                special_attr_power_map=SPECIAL_ATTR_POWER_MAP,
                total_attr_power_map=TOTAL_ATTR_POWER_MAP
            )
            
            # 进度回调
            def enum_progress_callback(current, total, message=""):
                if not self._cancel_flag.is_set():
                    progress = int(80 * current / total) if total > 0 else 0
                    self._compute_queue.put(("progress", progress, 100, message))
            
            # 枚举所有组合（一次性获得所有解）
            from math import comb   
            total_combinations = comb(len(filtered_gems), slots)
            
            # 根据是否有目标需求决定枚举策略
            effective_target_requirements = target_requirements if (target_requirements and len(target_requirements) > 0) else None

            if total_combinations > 50000:  # 大数据集使用内存安全方法
                if effective_target_requirements:
                    # 有目标需求：按需求评分排序
                    enum_solutions = opt.enum_all_combinations_ui_threaded_memory_safe_dual_topk(
                        filtered_gems,
                        slots=slots,
                        progress_callback=enum_progress_callback,
                        cancel_flag=self._cancel_flag,
                        max_keep=500,
                        target_requirements=effective_target_requirements
                    )
                else:
                    # 无目标需求：按战力分排序
                    enum_solutions = opt.enum_all_combinations_ui_threaded_memory_safe_dual_topk(
                        filtered_gems,
                        slots=slots,
                        progress_callback=enum_progress_callback,
                        cancel_flag=self._cancel_flag,
                        max_keep=500,
                        target_requirements=None
                    )
            else:
                if effective_target_requirements:
                    # 有目标需求：按需求评分排序
                    enum_solutions = opt.enum_all_combinations_ui_memory_safe_dual_topk(
                        filtered_gems,
                        slots=slots,
                        progress_callback=enum_progress_callback,
                        max_keep=500,
                        target_requirements=effective_target_requirements
                    )
                else:
                    # 无目标需求：按战力分排序
                    enum_solutions = opt.enum_all_combinations_ui_memory_safe_dual_topk(
                        filtered_gems,
                        slots=slots,
                        progress_callback=enum_progress_callback,
                        max_keep=500,
                        target_requirements=None
                    )

            if self._cancel_flag.is_set():
                return

            self._compute_queue.put(("progress", 80, 100, "生成不同排序方案..."))
            
            # 生成两种不同排序的方案
            if effective_target_requirements:
                # 有目标需求时：需求方案已经按需求评分排序
                req_list = enum_solutions[:30]
                
                # 战力方案需要单独按战力分重新排序
                battle_power_sorted = []
                for sol in enum_solutions:
                    bp = opt.combat_power_from_attr_breakdown(sol.attr_breakdown)
                    battle_power_sorted.append((bp, sol))
                battle_power_sorted.sort(key=lambda x: x[0], reverse=True)
                abs_list = [sol for bp, sol in battle_power_sorted[:30]]
            else:
                # 无目标需求时：enum_solutions已经按战力分排序
                abs_list = enum_solutions[:30]
                req_list = abs_list  # 需求方案等同于战力方案
            
            # 职业推荐（目前使用战力排序）
            bs_list = abs_list
            
            self._compute_queue.put(("progress", 100, 100, "完成"))
            self._compute_queue.put(("result", req_list, abs_list, bs_list, opt, target_requirements))

        except Exception as e:
            import traceback
            traceback.print_exc()
            self._compute_queue.put(("error", str(e)))

    def _poll_compute_progress_with_requirements(self):
        """轮询计算进度 - 支持需求方案的版本"""
        try:
            has_updates = False
            
            # 批量处理队列中的所有更新
            while True:
                try:
                    item = self._compute_queue.get_nowait()
                    msg_type = item[0]
                    
                    if msg_type == "progress":
                        _, current, total, message = item
                        self._progress_bar["maximum"] = max(1, total)
                        self._progress_bar["value"] = current
                        percentage = int(100 * current / total) if total > 0 else 0
                        self._progress_var.set(f"{percentage}% - {message}")
                        has_updates = True
                        
                    elif msg_type == "result":
                        _, req_list, abs_list, bs_list, opt, target_requirements = item
                        self._close_progress_window()
                        self._render_results_with_requirements(req_list, abs_list, bs_list, "（全部-多线程）", opt, target_requirements)
                        return
                        
                    elif msg_type == "error":
                        _, error_msg = item
                        self._close_progress_window()
                        messagebox.showerror("计算错误", f"多线程计算发生错误：{error_msg}")
                        return
                        
                except queue.Empty:
                    break
            
            # 只在有更新时刷新UI
            if has_updates:
                self.update_idletasks()
                
        except Exception as e:
            print(f"轮询进度时发生错误：{e}")
        
        # 继续轮询，1秒间隔
        if hasattr(self, '_progress_window') and self._progress_window.winfo_exists():
            self.after(1000, self._poll_compute_progress_with_requirements)

    def _cancel_computation(self):
        """取消计算"""
        if hasattr(self, '_cancel_flag'):
            self._cancel_flag.set()
        self._close_progress_window()

    def _close_progress_window(self):
        """关闭进度窗口"""
        try:
            if hasattr(self, '_progress_window') and self._progress_window.winfo_exists():
                self._progress_window.grab_release()
                self._progress_window.destroy()
        except:
            pass

    def _render_empty_results(self, suffix):
        """渲染空结果"""
        # 检查条件
        has_targets_or_excludes = False
        for ent, _ in self.targets:
            if ent.get_value():
                has_targets_or_excludes = True
                break
        if not has_targets_or_excludes:
            for ent in self.exclude_entries:
                if ent.get_value():
                    has_targets_or_excludes = True
                    break
        
        has_job = bool(self.job_field.get_value())
        
        # 最高等级方案始终显示空结果
        self.abs_text.configure(state="normal")
        self.abs_text.delete("1.0", tk.END)
        self.abs_text.insert(tk.END, f"计算完成{suffix}\n\n（无可用组合）")
        self.abs_text.configure(state="disabled")
        
        # 需求最佳方案根据条件显示
        self.req_text.configure(state="normal")
        self.req_text.delete("1.0", tk.END)
        if has_targets_or_excludes:
            self.req_text.insert(tk.END, f"计算完成{suffix}\n\n（无可用组合）")
        else:
            self.req_text.insert(tk.END, f"需求最佳方案{suffix}\n\n（需要设置期望属性或排除属性才会显示结果）")
        self.req_text.configure(state="disabled")
        
        # 战力分推荐方案根据条件显示
        self.bs_text.configure(state="normal")
        self.bs_text.delete("1.0", tk.END)
        if has_job:
            self.bs_text.insert(tk.END, f"计算完成{suffix}\n\n（无可用组合）")
        else:
            self.bs_text.insert(tk.END, f"战力分推荐方案{suffix}\n\n（需要选择职业才会显示结果）")
        self.bs_text.configure(state="disabled")

    def _render_results_with_requirements(self, req_list, abs_list, bs_list, suffix, opt, target_requirements):
        """渲染计算结果 - 增强版，显示目标需求信息"""
        # 保存结果用于导出
        self._last_req_list = list(req_list)
        self._last_abs_list = list(abs_list)
        self._last_bs_list = list(bs_list)

        def combo_lines(sol): 
            # 兼容处理：无论是ModuleSolution对象还是元组都能正确获取modules
            if hasattr(sol, 'modules'):
                mods = sol.modules
            elif isinstance(sol, (list, tuple)) and len(sol) > 2:
                # 旧格式：(score, something, modules, ...)
                mods = sol[2] if len(sol) > 2 else sol
            else:
                mods = sol
            return "\n".join([format_gem_str(g) for g in mods])
        
        def get_attr_breakdown(sol):
            # 兼容处理：获取属性分布
            if hasattr(sol, 'attr_breakdown'):
                return sol.attr_breakdown
            elif isinstance(sol, (list, tuple)) and len(sol) > 3:
                # 旧格式：(score, something, modules, attr_breakdown, ...)
                return sol[3] if len(sol) > 3 else {}
            else:
                # 如果无法获取，尝试从模块重新计算
                try:
                    mods = combo_lines(sol).split('\n')  # 这里需要更好的解析
                    # 这里应该调用opt的计算方法重新计算attr_breakdown
                    return {}
                except:
                    return {}
        
        def render_one(box, title, data, show_box=True, show_req_score=False):
            box.configure(state="normal")
            box.delete("1.0", tk.END)
            
            if not show_box:
                box.insert(tk.END, f"{title}\n\n（需要设置相应条件才会显示结果）")
                box.configure(state="disabled")
                return
                
            box.insert(tk.END, f"{title}\n")
            
            # 显示目标需求信息
            if target_requirements and show_req_score:
                req_text = "目标需求: " + ", ".join([f"{attr}≥{level}级" for attr, level in target_requirements])
                box.insert(tk.END, f"{req_text}\n")
            
            box.insert(tk.END, "\n")
            
            # 限制显示30个方案
            display_data = data[:30]
            
            for k, sol in enumerate(display_data, start=1):
                attr_breakdown = get_attr_breakdown(sol)
                levels = _levels_from_points(attr_breakdown)
                
                # 计算战力分
                try:
                    if hasattr(opt, 'combat_power_from_attr_breakdown'):
                        bp = opt.combat_power_from_attr_breakdown(attr_breakdown)
                    else:
                        bp = sum(levels.values())
                except Exception:
                    bp = sum(levels.values())
                
                # 如果是需求方案，还要显示需求评分
                req_score_text = ""
                if show_req_score and target_requirements:
                    try:
                        req_score = opt.requirement_score_from_attr_breakdown(attr_breakdown, target_requirements)
                        req_score_text = f"需求匹配分数：{req_score:.0f}\n"
                    except Exception:
                        pass
                        
                # 获取模块信息
                try:
                    modules_text = combo_lines(sol)
                except Exception:
                    modules_text = "（模块信息解析失败）"
                    
                box.insert(tk.END, f"方案 #{k}\n")
                box.insert(tk.END, f"{modules_text}\n")
                box.insert(tk.END, f"等级分布：\n{_pretty_levels_table(levels)}\n")
                box.insert(tk.END, f"战斗力分数：{bp}\n")
                if req_score_text:
                    box.insert(tk.END, req_score_text)
                box.insert(tk.END, "\n")
                    
            if not display_data:
                box.insert(tk.END, "（无）\n")
            elif len(data) > 30:
                box.insert(tk.END, f"（仅显示前30个方案，共{len(data)}个）\n")
                    
            box.configure(state="disabled")

        # 检查是否有目标属性或排除属性
        has_targets_or_excludes = bool(target_requirements and len(target_requirements) > 0) or any(ent.get_value() for ent in self.exclude_entries)
        
        # 检查是否选择了职业
        has_job = bool(self.job_field.get_value())

        # 按新顺序渲染
        render_one(self.abs_text, f"战斗力最高方案{suffix}（按战斗力降序）", abs_list, True, False)
        render_one(self.req_text, f"需求最佳方案{suffix}（按需求匹配降序）", req_list, has_targets_or_excludes, bool(target_requirements and len(target_requirements) > 0))
        render_one(self.bs_text, f"职业推荐方案{suffix}（按战斗力降序）", bs_list, has_job, False)

class MultiCheckDialog(tk.Toplevel):
    """
    简单多选弹窗：
    - 每行4个候选项（Checkbutton）
    - 无滚动条、无搜索框、无“无”选项
    - 限制最大选择数量（max_select）
    - 返回 self.result = [选中的字符串...]
    """
    def __init__(self, master, title, options, preselected=None, max_select=1, columns=4):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.result = None
        self.max_select = int(max_select)
        self.vars = []
        self.opts = list(options)
        self.columns = max(1, int(columns))

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        # 计数标签
        self.var_cnt = tk.StringVar(value=f"已选 0 / {self.max_select}")
        ttk.Label(frm, textvariable=self.var_cnt).grid(row=0, column=0, columnspan=self.columns, sticky="w", pady=(0,6))

        # 网格勾选
        row = 1
        col = 0
        pre = set(preselected or [])
        for name in self.opts:
            v = tk.BooleanVar(value=(name in pre))
            cb = ttk.Checkbutton(frm, text=name, variable=v, command=self._on_toggle)
            cb.grid(row=row, column=col, padx=6, pady=6, sticky="w")
            self.vars.append((name, v))
            col += 1
            if col >= self.columns:
                col = 0
                row += 1

        # 底部按钮
        btnf = ttk.Frame(frm)
        btnf.grid(row=row+1, column=0, columnspan=self.columns, pady=(8,0), sticky="e")
        ttk.Button(btnf, text="清空", command=self._clear).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(btnf, text="确定", command=self._ok).pack(side=tk.LEFT, padx=(0,8))
        ttk.Button(btnf, text="取消", command=self._cancel).pack(side=tk.LEFT)

        # 初始化计数
        self._update_count()

        # 居中
        self.update_idletasks()
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

        self.bind("<Escape>", lambda e: self._cancel())

    def _sel_names(self):
        return [name for name, v in self.vars if v.get()]

    def _update_count(self):
        n = len(self._sel_names())
        self.var_cnt.set(f"已选 {n} / {self.max_select}")

    def _on_toggle(self):
        names = self._sel_names()
        if len(names) > self.max_select:
            # 超出则把刚才最后一个勾选强制退回
            # 简单做法：把所有选中的最后一个取消
            last = names[-1]
            for name, v in self.vars:
                if name == last:
                    v.set(False)
                    break
        self._update_count()

    def _ok(self):
        self.result = self._sel_names()
        self.destroy()

    def _clear(self):
        for _, v in self.vars:
            v.set(False)
        self._update_count()

    def _cancel(self):
        self.result = None
        self.destroy()


#== 选择控件们（保留不变）==
class SimpleListDialog(tk.Toplevel):
    """简单单选列表弹窗：支持搜索、单选、双击选择；含“（无）”选项。"""
    def __init__(self, master, title, options, current=None):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.result = None
        # 在最前加一个“（无）”
        self._opts_all = ["（无）"] + list(options)
        self.bind("<Escape>", lambda e: self._cancel())
        self.protocol("WM_DELETE_WINDOW", self._cancel)

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        top = ttk.Frame(frm)
        top.pack(fill=tk.X)
        ttk.Label(top, text="搜索：").pack(side=tk.LEFT)
        self.var_q = tk.StringVar()
        ent = ttk.Entry(top, textvariable=self.var_q, width=24)
        ent.pack(side=tk.LEFT, padx=(6,0))
        ent.bind("<KeyRelease>", self._on_search)

        self.lb = tk.Listbox(frm, height=12, exportselection=False)
        self.lb.pack(fill=tk.BOTH, expand=True, pady=8)
        self.lb.bind("<Double-Button-1>", lambda e: self._confirm())
        self._reload_list()

        if current in self._opts_all:
            idx = self._opts_all.index(current)
            self.lb.selection_set(idx)
            self.lb.see(idx)

        self.update_idletasks()
        try:
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception:
            pass

    def _reload_list(self, filtered=None):
        data = filtered if filtered is not None else self._opts_all
        self.lb.delete(0, tk.END)
        for it in data:
            self.lb.insert(tk.END, it)

    def _on_search(self, *_):
        q = self.var_q.get().strip().lower()
        if not q:
            self._reload_list(); return
        self._reload_list([o for o in self._opts_all if q in o.lower()])

    def _confirm(self):
        sel = self.lb.curselection()
        if not sel: return
        val = self.lb.get(sel[0])
        self.result = None if val == "（无）" else val
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()


class AttributePicker(tk.Toplevel):
    """
    选择词条弹窗：搜索 + 列表；仅双击选择；含“（无）”选项。
    on_ok(name|None)
    """
    def __init__(self, master, all_attrs, disabled_set=None, title="选择词条", on_ok=None, init_text=""):
        super().__init__(master)
        self.title(title)
        self.resizable(False, False)
        self.grab_set()
        self.bind("<Escape>", lambda e: self._on_cancel())
        self.protocol("WM_DELETE_WINDOW", self._on_cancel)

        self.all_attrs = ["（无）"] + list(all_attrs)
        self.disabled = set(disabled_set or [])
        self.on_ok = on_ok

        frm = ttk.Frame(self, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        top = ttk.Frame(frm)
        top.pack(fill=tk.X)
        ttk.Label(top, text="搜索：").pack(side=tk.LEFT)
        self.var_q = tk.StringVar(value=init_text)
        ent = ttk.Entry(top, textvariable=self.var_q, width=28)
        ent.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(4,0))
        ent.bind("<KeyRelease>", self._on_filter)

        mid = ttk.Frame(frm)
        mid.pack(fill=tk.BOTH, expand=True, pady=(8,0))
        self.lb = tk.Listbox(mid, height=12)
        self.lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self.lb.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.lb.configure(yscrollcommand=sb.set)

        self.lb.bind("<Double-Button-1>", self._on_dbl)
        self._reload_list()
        self._center_to(master)

    def _center_to(self, master):
        try:
            self.update_idletasks()
            x = master.winfo_rootx() + (master.winfo_width()-self.winfo_width())//2
            y = master.winfo_rooty() + (master.winfo_height()-self.winfo_height())//2
            self.geometry(f"+{max(0,x)}+{max(0,y)}")
        except Exception: pass

    def _on_filter(self, e=None):
        self._reload_list()

    def _reload_list(self):
        q = self.var_q.get().strip().lower()
        cand = [a for a in self.all_attrs if q in a.lower()] if q else list(self.all_attrs)
        cand.sort(key=lambda x: "0" if x=="（无）" else x)
        self.lb.delete(0, tk.END)
        for a in cand:
            if a!="（无）" and a in self.disabled:
                self.lb.insert(tk.END, f"× {a}")
            else:
                self.lb.insert(tk.END, a)

    def _pick_value(self):
        cur = self.lb.curselection()
        if not cur: return None
        val = self.lb.get(cur[0])
        if val == "（无）": return None
        if val.startswith("× "):
            val = val[2:]
            messagebox.showwarning("不可选择", f"“{val}”已被占用。")
            return None
        return val

    def _on_dbl(self, e=None):
        v = self._pick_value()
        if callable(self.on_ok): self.on_ok(v)
        self.destroy()

    def _on_cancel(self):
        if callable(self.on_ok): self.on_ok(None)
        self.destroy()


class ClickSelectEntry(ttk.Frame):
    """
    无下拉、无按钮：只读输入框，点击后弹出选择窗口。
    """
    def __init__(self, master, options, placeholder="", width=16, on_select=None):
        super().__init__(master)
        self.options = list(options)
        self.placeholder = placeholder
        self.on_select = on_select
        self._popup = None  # 防止重复弹出

        self.var = tk.StringVar(value=placeholder)
        self.entry = ttk.Entry(self, textvariable=self.var, width=width, state="readonly")
        self.entry.pack(fill="x", expand=True)
        self.entry.bind("<Button-1>", self._open_popup)

    def get(self):
        v = (self.var.get() or "").strip()
        return "" if v == self.placeholder else v

    def set(self, v: str):
        self.var.set(v or self.placeholder)

    def _open_popup(self, event=None):
        if not self.options or (self._popup and self._popup.winfo_exists()):
            return
        top = tk.Toplevel(self)
        self._popup = top
        top.title("请选择")
        top.transient(self.winfo_toplevel())
        top.grab_set()
        top.resizable(False, False)
        top.bind("<Escape>", lambda e: self._close_popup(top))
        top.protocol("WM_DELETE_WINDOW", lambda: self._close_popup(top))

        frm = ttk.Frame(top, padding=8)
        frm.pack(fill="both", expand=True)

        sv = tk.StringVar()
        ent_search = ttk.Entry(frm, textvariable=sv)
        ent_search.pack(fill="x", pady=(0,6))
        ent_search.focus_set()

        lb = tk.Listbox(frm, height=12)
        lb.pack(fill="both", expand=True)

        opts = ["（无）"] + [o for o in self.options if o != "（无）" and o.strip()]

        def _refresh():
            kw = sv.get().strip().lower()
            lb.delete(0, tk.END)
            for it in opts:
                if not kw or kw in it.lower():
                    lb.insert(tk.END, it)
        _refresh()

        def _ok():
            sel = lb.curselection()
            if sel:
                val = lb.get(sel[0])
                val = None if val == "（无）" else val
                self.var.set("" if val is None else val)
                if callable(self.on_select):
                    self.on_select(val)
            self._close_popup(top)

        lb.bind("<Double-Button-1>", lambda e: _ok())
        sv.trace_add("write", lambda *_: _refresh())

        top.update_idletasks()
        try:
            x = self.winfo_rootx() + (self.winfo_width()-top.winfo_width())//2
            y = self.winfo_rooty() + self.winfo_height() + 6
            top.geometry(f"+{x}+{y}")
        except Exception:
            pass


    def _close_popup(self, top):
        try:
            top.grab_release()
        except Exception:
            pass
        try:
            top.destroy()
        except Exception:
            pass
        self._popup = None


#== 组合控件：只读显示=========
class AttrPickerField(ttk.Frame):
    """
    只读 Entry，点击弹出 AttributePicker。
    - 支持 set_disabled()/get_value()/var，兼容原互斥逻辑
    - on_change 回调：值变化时触发
    """
    def __init__(self, master, values, on_change=None, width=24, placeholder="", open_handler=None):
        super().__init__(master)
        self.values = list(values)
        self.disabled = set()
        self.on_change = on_change
        self.placeholder = placeholder or ""
        self._open_handler = open_handler

        self.var = tk.StringVar(value=self.placeholder)
        self.ent = ttk.Entry(self, textvariable=self.var, width=width, state="readonly")
        self.ent.pack(side=tk.LEFT, fill="x", expand=True)

        if callable(self._open_handler):
            self.ent.bind("<Button-1>", self._open_handler)
        else:
            self.ent.bind("<Button-1>", self._open_picker)

    def _open_picker(self, *_):
        init = self.get_value() or self.placeholder
        AttributePicker(
            self.winfo_toplevel(),
            self.values,
            disabled_set=self.disabled,
            title="选择词条",
            on_ok=self._on_selected,
            init_text=init
        )

    def _on_selected(self, name):
        if name is None:
            return
        self.var.set(name)
        if callable(self.on_change):
            self.on_change(self)

    def set_disabled(self, names:set):
        self.disabled = set(names or [])
        cur = self.get_value()
        if cur and cur in self.disabled:
            self.var.set(self.placeholder)
            if callable(self.on_change):
                self.on_change(self)

    def get_value(self):
        v = (self.var.get() or "").strip()
        return "" if v == self.placeholder else v


if __name__ == "__main__":
    App().mainloop()
